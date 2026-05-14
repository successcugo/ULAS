"""
advisor_app.py — ULAS Course Advisor Portal

Access levels:
  ICT (master)  — Can create/delete advisor accounts for any dept.
                  Can change any advisor's password.
                  Credentials set in Streamlit secrets as ICT_USERNAME / ICT_PASSWORD.
  Advisor       — Can manage reps in their own department only.
                  Can change passwords of reps AND co-advisors in their own dept.
                  Cannot create or delete advisor accounts.
"""
from __future__ import annotations

import streamlit as st
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))

from futo_data import get_schools, get_departments, get_levels, get_full_structure, save_structure, invalidate_structure_cache, get_school_abbr
from core import (
    load_active_semester, start_semester, end_semester, load_semester_history,
    load_student_gpa, assign_semester_gpa, assign_gpa_for_semester,
    compute_cgpa, get_dept_matric_numbers, get_dept_students,
    lava_sem_path, get_available_sessions, get_semesters_for_session,
    authenticate_user, authenticate_ict, load_users, create_user,
    update_password, delete_user, get_reps_for_dept, get_advisors_for_dept,
    load_settings, save_settings,
    load_advisor_lifetime, save_advisor_lifetime,
    get_all_advisors, load_settings, save_settings, hash_password,
    verify_password, futo_now_str, get_dept_abbreviation, set_dept_abbreviation,
)
from github_store import invalidate_cache, upload_chat_file
from chat_store import (
    load_room, post_message, delete_message,
    get_unread, mark_read, count_unread,
    build_display_name, all_rooms_for_advisor,
    dm_room, school_room,
)

st.set_page_config(
    page_title="ULAS Advisor Portal",
    page_icon="🧑‍🏫",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;900&display=swap');
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
.hero {
    background: linear-gradient(135deg, #003087 0%, #1a56db 100%);
    border-radius: 16px; padding: 2rem; text-align: center;
    color: white; margin-bottom: 2rem;
    box-shadow: 0 4px 20px rgba(0,48,135,0.3);
}
.hero h1 { font-size: 2rem; font-weight: 900; margin: 0; }
.hero .sub { opacity: 0.85; margin-top: 0.3rem; font-size: 0.9rem; }
.hero-ict {
    background: linear-gradient(135deg, #5c0099 0%, #9b30ff 100%);
    border-radius: 16px; padding: 2rem; text-align: center;
    color: white; margin-bottom: 2rem;
    box-shadow: 0 4px 20px rgba(92,0,153,0.3);
}
.hero-ict h1 { font-size: 2rem; font-weight: 900; margin: 0; }
.hero-ict .sub { opacity: 0.85; margin-top: 0.3rem; font-size: 0.9rem; }
.info-card {
    background: rgba(26, 86, 219, 0.08);
    border-left: 4px solid #1a56db;
    border-radius: 0 8px 8px 0;
    padding: 0.7rem 1rem; margin: 0.5rem 0; color: inherit;
}
.info-card b { color: #1a56db; }
.ict-card {
    background: rgba(155, 48, 255, 0.08);
    border-left: 4px solid #9b30ff;
    border-radius: 0 8px 8px 0;
    padding: 0.7rem 1rem; margin: 0.5rem 0; color: inherit;
}
.ict-card b { color: #9b30ff; }
.rep-card {
    background: rgba(26, 86, 219, 0.05);
    border: 1px solid rgba(26, 86, 219, 0.2);
    border-radius: 10px; padding: 0.9rem 1.1rem; margin-bottom: 0.7rem;
    color: inherit;
}
.adv-card {
    background: rgba(155, 48, 255, 0.05);
    border: 1px solid rgba(155, 48, 255, 0.2);
    border-radius: 10px; padding: 0.9rem 1.1rem; margin-bottom: 0.7rem;
    color: inherit;
}
.badge-rep {
    background: rgba(26, 86, 219, 0.15); color: #1a56db;
    border-radius: 20px; padding: 3px 12px; font-size: 0.8rem; font-weight: 700;
}
.badge-adv {
    background: rgba(155, 48, 255, 0.15); color: #9b30ff;
    border-radius: 20px; padding: 3px 12px; font-size: 0.8rem; font-weight: 700;
}
div[data-testid="stForm"] {
    border: 1.5px solid rgba(128,128,128,0.3);
    border-radius: 12px; padding: 1.2rem 1.2rem 0.5rem;
}
.stButton > button { border-radius: 8px; font-weight: 600; }
</style>
""", unsafe_allow_html=True)

# ── Fixed footer ─────────────────────────────────────────────────────────────
st.markdown('\n<style>\n.ulas-footer {\n    position: fixed;\n    bottom: 0; left: 0; right: 0;\n    text-align: center;\n    padding: 0.45rem 1rem;\n    font-size: 0.78rem;\n    background: rgba(0,0,0,0.45);\n    backdrop-filter: blur(6px);\n    -webkit-backdrop-filter: blur(6px);\n    color: rgba(255,255,255,0.55);\n    letter-spacing: 0.04em;\n    z-index: 9999;\n    border-top: 1px solid rgba(255,255,255,0.07);\n}\n.ulas-footer b { color: rgba(255,255,255,0.8); font-weight: 600; }\n.ulas-footer .dot { color: rgba(255,255,255,0.3); margin: 0 0.3em; }\n/* Push content up so footer never overlaps last element */\nsection.main > div { padding-bottom: 2.8rem !important; }\n</style>\n<div class="ulas-footer">\n    Made with ❤️ by\n    <b>SESET</b><span class="dot">•</span><b>EPE</b><span class="dot">•</span><b>2030/2031</b>\n</div>\n', unsafe_allow_html=True)

st.markdown("""
<div class="hero">
    <h1>🧑‍🏫 ULAS Advisor Portal</h1>
    <div class="sub">Course Advisor Management Console — FUTO</div>
</div>
""", unsafe_allow_html=True)


# ── Session state ─────────────────────────────────────────────────────────────
for k, v in {
    "portal_role": None,   # 'ict' | 'advisor'
    "adv_user": None,      # advisor user dict (if role == 'advisor')
    # Cascading dropdowns
    "adv_na_school": None, "adv_na_dept": None,
    "bs_school": None,     "bs_dept": None,
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

def do_logout():
    st.session_state.portal_role  = None
    st.session_state.adv_user     = None
    st.session_state.adv_na_school = None
    st.session_state.adv_na_dept   = None
    invalidate_cache("__users")
    invalidate_cache("__settings")


# ── Cascading dropdown helpers ─────────────────────────────────────────────────
def _on_na_school():
    st.session_state.adv_na_school = st.session_state._adv_na_school_w
    st.session_state.adv_na_dept   = None

def _on_na_dept():
    st.session_state.adv_na_dept = st.session_state._adv_na_dept_w

def _on_bs_school():
    st.session_state.bs_school = st.session_state._bs_school_w
    st.session_state.bs_dept   = None

def _on_bs_dept():
    st.session_state.bs_dept = st.session_state._bs_dept_w



# ── Chat session-state defaults ─────────────────────────────────────────────
for _ck in ["chat_room", "chat_dm_target", "chat_poll_ts", "chat_messages_cache",
            "chat_unread_cache", "chat_file_upload"]:
    if _ck not in st.session_state:
        st.session_state[_ck] = None

# ── Branded error screen ─────────────────────────────────────────────────────
def _show_error(exc: Exception):
    st.markdown("""
    <div style="
        max-width:520px; margin:4rem auto; padding:2.5rem 2rem;
        background:rgba(123,0,0,0.08); border:1.5px solid rgba(123,0,0,0.25);
        border-radius:16px; text-align:center;
    ">
        <div style="font-size:3rem; margin-bottom:0.5rem;">⚠️</div>
        <h2 style="color:#c0392b; margin:0 0 0.5rem;">Something went wrong</h2>
        <p style="opacity:0.7; font-size:0.92rem; margin:0 0 1.5rem;">
            An unexpected error occurred. Please try refreshing the page.<br>
            If the problem persists, contact ICT administration.
        </p>
        <button onclick="window.location.reload()"
            style="background:#c0392b; color:white; border:none; border-radius:8px;
                   padding:0.6rem 1.8rem; font-size:1rem; cursor:pointer; font-weight:600;">
            🔄 Refresh Page
        </button>
    </div>
    """, unsafe_allow_html=True)

# ── Main app body ─────────────────────────────────────────────────────────────
try:
    # ═══════════════════════════════════════════════════════════════════════════════
    #  LOGIN SCREEN
    # ═══════════════════════════════════════════════════════════════════════════════
    if st.session_state.portal_role is None:
        st.markdown("## Login")
        st.markdown("Log in as a **Course Advisor** or as **ICT** (master admin).")

        login_tab, ict_tab = st.tabs(["🧑‍🏫 Course Advisor Login", "🔐 ICT Master Login"])

        with login_tab:
            with st.form("adv_login"):
                uname   = st.text_input("Advisor Username")
                pwd     = st.text_input("Password", type="password")
                log_btn = st.form_submit_button("Login as Advisor", type="primary")
            if log_btn:
                with st.spinner("Authenticating..."):
                    user = authenticate_user(uname, pwd, role="advisor")
                if user:
                    st.session_state.portal_role = "advisor"
                    st.session_state.adv_user    = user
                    st.rerun()
                else:
                    st.error("Invalid advisor credentials.")

        with ict_tab:
            with st.form("ict_login"):
                ict_u   = st.text_input("ICT Username")
                ict_p   = st.text_input("ICT Password", type="password")
                ict_btn = st.form_submit_button("Login as ICT", type="primary")
            if ict_btn:
                if authenticate_ict(
                    ict_u, ict_p,
                    expected_user=st.secrets.get("ICT_USERNAME", ""),
                    expected_pw=st.secrets.get("ICT_PASSWORD", ""),
                ):
                    st.session_state.portal_role = "ict"
                    st.rerun()
                else:
                    st.error("Invalid ICT credentials.")

        # Bootstrap: if no advisors yet, only show ICT login and a note
        if not get_all_advisors():
            st.info("ℹ️ No advisor accounts exist yet. Log in as ICT to create the first advisor.")

        st.stop()


    # ═══════════════════════════════════════════════════════════════════════════════
    #  ICT MASTER DASHBOARD
    # ═══════════════════════════════════════════════════════════════════════════════
    if st.session_state.portal_role == "ict":

        st.markdown("""
        <div class="hero-ict">
            <h1>🔐 ICT Master Console</h1>
            <div class="sub">Full administrative access — handle with care</div>
        </div>
        """, unsafe_allow_html=True)

        hc, lc = st.columns([5, 1])
        with hc:
            st.markdown('<div class="ict-card"><b>Logged in as:</b> ICT Master</div>',
                        unsafe_allow_html=True)
        with lc:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Logout", key="ict_logout"):
                do_logout()
                st.rerun()

        ict_tab1, ict_tab2, ict_tab3, ict_tab4, ict_tab5 = st.tabs(
            ["👥 All Advisors", "➕ Create Advisor", "🏫 Schools & Depts", "⚙️ Settings", "📅 Semester"]
        )

        # ── ICT TAB 1 — View & manage all advisors ────────────────────────────────
        with ict_tab1:
            st.markdown("### All Advisor Accounts")
            invalidate_cache("__users")
            advisors = get_all_advisors()
            if not advisors:
                st.info("No advisor accounts exist yet. Create one in the next tab.")
            else:
                for adv in sorted(advisors, key=lambda a: (a["school"], a["department"], a["username"])):
                    ac, pc, dc = st.columns([5, 2, 1])
                    with ac:
                        st.markdown(f"""
                        <div class="adv-card">
                            <span class="badge-adv">Advisor</span>
                            &nbsp;&nbsp;<b>{adv['username']}</b>
                            &nbsp;·&nbsp; {adv['department']}
                            <span style="font-size:0.8rem;opacity:0.6">
                                &nbsp;·&nbsp; {adv['school'][:30]}
                            </span>
                        </div>""", unsafe_allow_html=True)
                    with pc:
                        if st.button("🔑 Reset PW", key=f"ict_pw_{adv['username']}"):
                            st.session_state[f"ict_reset_{adv['username']}"] = True
                    with dc:
                        if st.button("🗑️", key=f"ict_del_{adv['username']}", help=f"Delete {adv['username']}"):
                            st.session_state[f"ict_delconfirm_{adv['username']}"] = True

                    # Inline password reset
                    if st.session_state.get(f"ict_reset_{adv['username']}"):
                        with st.form(f"ict_pw_form_{adv['username']}"):
                            np  = st.text_input("New Password", type="password", key=f"np_{adv['username']}")
                            np2 = st.text_input("Confirm",     type="password", key=f"np2_{adv['username']}")
                            ok_btn  = st.form_submit_button("Set Password", type="primary")
                            can_btn = st.form_submit_button("Cancel")
                        if ok_btn:
                            if np != np2:
                                st.error("Passwords do not match.")
                            elif len(np) < 6:
                                st.error("Minimum 6 characters.")
                            else:
                                update_password(adv["username"], np)
                                invalidate_cache("__users")
                                st.success(f"Password updated for {adv['username']}.")
                                st.session_state.pop(f"ict_reset_{adv['username']}", None)
                                st.rerun()
                        if can_btn:
                            st.session_state.pop(f"ict_reset_{adv['username']}", None)
                            st.rerun()

                    # Inline delete confirm
                    if st.session_state.get(f"ict_delconfirm_{adv['username']}"):
                        st.warning(f"Delete advisor **{adv['username']}** ({adv['department']})? This cannot be undone.")
                        y, n = st.columns(2)
                        with y:
                            if st.button("Yes, delete", type="primary", key=f"ict_delyes_{adv['username']}"):
                                delete_user(adv["username"])
                                invalidate_cache("__users")
                                st.session_state.pop(f"ict_delconfirm_{adv['username']}", None)
                                st.rerun()
                        with n:
                            if st.button("Cancel", key=f"ict_delno_{adv['username']}"):
                                st.session_state.pop(f"ict_delconfirm_{adv['username']}", None)
                                st.rerun()

        # ── ICT TAB 2 — Create advisor ────────────────────────────────────────────
        with ict_tab2:
            st.markdown("### Create New Advisor Account")
            st.info("Only ICT can create advisor accounts. Advisors are then responsible for their own departments.")

            schools   = get_schools()
            na_s_opts = ["— select school —"] + schools
            cur_nas   = st.session_state.adv_na_school
            st.selectbox(
                "School", na_s_opts,
                index=na_s_opts.index(cur_nas) if cur_nas in na_s_opts else 0,
                key="_adv_na_school_w", on_change=_on_na_school,
            )

            na_depts  = get_departments(st.session_state.adv_na_school) if st.session_state.adv_na_school else []
            na_d_opts = ["— select department —"] + na_depts
            cur_nad   = st.session_state.adv_na_dept
            st.selectbox(
                "Department", na_d_opts,
                index=na_d_opts.index(cur_nad) if cur_nad in na_d_opts else 0,
                key="_adv_na_dept_w", on_change=_on_na_dept,
                disabled=not st.session_state.adv_na_school,
            )

            with st.form("ict_create_adv"):
                na_uname = st.text_input("Advisor Username")
                na_pwd   = st.text_input("Password", type="password")
                na_pwd2  = st.text_input("Confirm Password", type="password")
                na_btn   = st.form_submit_button("Create Advisor", type="primary")

            if na_btn:
                s = st.session_state.adv_na_school
                d = st.session_state.adv_na_dept
                if not s or not d:
                    st.error("Please select school and department above.")
                elif not na_uname.strip() or not na_pwd.strip():
                    st.error("All fields required.")
                elif na_pwd != na_pwd2:
                    st.error("Passwords do not match.")
                elif len(na_pwd) < 6:
                    st.error("Password must be at least 6 characters.")
                else:
                    ok, msg = create_user(na_uname.strip(), na_pwd, "advisor", s, d, None, "ICT")
                    if ok:
                        invalidate_cache("__users")
                        st.success(f"✅ Advisor '{na_uname}' created for {d}.")
                        st.session_state.adv_na_school = None
                        st.session_state.adv_na_dept   = None
                    else:
                        st.error(msg)

        # ── ICT TAB 3 — Schools & Departments ───────────────────────────────────────
        with ict_tab3:
            st.markdown("### 🏫 Schools & Departments")
            st.caption(
                "Changes take effect immediately across all ULAS apps. "
                "Deleting a school or department does **not** remove existing user accounts or session data."
            )

            invalidate_structure_cache()
            struct    = get_full_structure()
            schools_d = struct.get("schools", {})
            abbrevs   = struct.get("abbreviations", {})

            def _save(s, a):
                return save_structure({"schools": s, "abbreviations": a})

            # ── Add School ────────────────────────────────────────────────────────
            with st.expander("➕ Add New School", expanded=False):
                with st.form("add_school_form"):
                    ns_name  = st.text_input("Full school name", placeholder="School of XYZ (SXYZ)")
                    ns_abbr  = st.text_input("School abbreviation", placeholder="SXYZ", max_chars=8)
                    ns_btn   = st.form_submit_button("Add School", type="primary")
                if ns_btn:
                    if not ns_name.strip() or not ns_abbr.strip():
                        st.error("Both name and abbreviation are required.")
                    elif ns_name.strip() in schools_d:
                        st.error("A school with that name already exists.")
                    else:
                        schools_d[ns_name.strip()] = {}
                        abbrevs[ns_name.strip()]   = ns_abbr.strip().upper()
                        if _save(schools_d, abbrevs):
                            st.success(f"✅ School '{ns_name.strip()}' added.")
                            st.rerun()
                        else:
                            st.error("GitHub write failed.")

            st.divider()

            # ── Per-school management ─────────────────────────────────────────────
            if not schools_d:
                st.info("No schools defined yet. Add one above.")
            else:
                sel_school = st.selectbox(
                    "Select school to manage",
                    sorted(schools_d.keys()),
                    key="ict_sel_school",
                )
                depts = schools_d.get(sel_school, {})
                cur_abbr = abbrevs.get(sel_school, "")

                sc1, sc2 = st.columns([3, 1])
                with sc1:
                    st.markdown(f"""<div class="ict-card">
                        <b>{sel_school}</b> &nbsp;·&nbsp;
                        Abbreviation: <b>{cur_abbr}</b> &nbsp;·&nbsp;
                        {len(depts)} department(s)
                    </div>""", unsafe_allow_html=True)
                with sc2:
                    if st.button("🗑️ Delete School", key="del_school_btn"):
                        st.session_state["confirm_del_school"] = sel_school

                # Delete school confirm
                if st.session_state.get("confirm_del_school") == sel_school:
                    st.warning(
                        f"Delete **{sel_school}** and all its departments? "
                        "Existing user accounts will still exist but their school will no longer appear in dropdowns."
                    )
                    cy, cn = st.columns(2)
                    with cy:
                        if st.button("Yes, delete school", type="primary", key="yes_del_school"):
                            del schools_d[sel_school]
                            abbrevs.pop(sel_school, None)
                            _save(schools_d, abbrevs)
                            st.session_state.pop("confirm_del_school", None)
                            st.rerun()
                    with cn:
                        if st.button("Cancel", key="no_del_school"):
                            st.session_state.pop("confirm_del_school", None)
                            st.rerun()

                # Edit school abbreviation inline
                with st.expander("✏️ Edit School Abbreviation"):
                    with st.form("edit_abbr_form"):
                        new_abbr = st.text_input("Abbreviation", value=cur_abbr, max_chars=8)
                        ea_btn   = st.form_submit_button("Save Abbreviation", type="primary")
                    if ea_btn:
                        if new_abbr.strip():
                            abbrevs[sel_school] = new_abbr.strip().upper()
                            if _save(schools_d, abbrevs):
                                st.success(f"Abbreviation updated to {new_abbr.strip().upper()}.")
                                st.rerun()
                            else:
                                st.error("GitHub write failed.")

                st.markdown(f"#### Departments in {abbrevs.get(sel_school, sel_school)}")

                # ── Add Department ────────────────────────────────────────────────
                with st.expander("➕ Add Department to this School"):
                    with st.form("add_dept_form"):
                        nd_name   = st.text_input("Department name", placeholder="e.g. Software Engineering")
                        nd_levels = st.number_input("Number of levels", min_value=1, max_value=8, value=4, step=1)
                        nd_btn    = st.form_submit_button("Add Department", type="primary")
                    if nd_btn:
                        if not nd_name.strip():
                            st.error("Department name cannot be empty.")
                        elif nd_name.strip() in depts:
                            st.error("That department already exists in this school.")
                        else:
                            schools_d[sel_school][nd_name.strip()] = int(nd_levels)
                            if _save(schools_d, abbrevs):
                                st.success(f"✅ Department '{nd_name.strip()}' added with {nd_levels} levels.")
                                st.rerun()
                            else:
                                st.error("GitHub write failed.")

                # ── List & manage existing departments ────────────────────────────
                if not depts:
                    st.info("No departments in this school yet.")
                else:
                    for dept_name, num_levels in sorted(depts.items()):
                        dc1, dc2, dc3 = st.columns([5, 2, 1])
                        safe_key = dept_name.replace(" ", "_").replace("(", "").replace(")", "")

                        with dc1:
                            st.markdown(f"""<div class="rep-card">
                                <b>{dept_name}</b>
                                &nbsp;·&nbsp;
                                <span style="opacity:0.65">{num_levels} levels
                                ({", ".join(str((i+1)*100) for i in range(int(num_levels)))})</span>
                            </div>""", unsafe_allow_html=True)

                        with dc2:
                            # Inline level editor
                            if st.button("✏️ Edit Levels", key=f"edit_lvl_{safe_key}"):
                                st.session_state[f"editing_{safe_key}"] = True

                        with dc3:
                            if st.button("🗑️", key=f"del_dept_{safe_key}", help=f"Delete {dept_name}"):
                                st.session_state[f"confirm_del_dept_{safe_key}"] = True

                        # Delete department confirm
                        if st.session_state.get(f"confirm_del_dept_{safe_key}"):
                            st.warning(f"Delete department **{dept_name}**?")
                            yy, nn = st.columns(2)
                            with yy:
                                if st.button("Yes, delete", type="primary", key=f"yes_del_dept_{safe_key}"):
                                    del schools_d[sel_school][dept_name]
                                    _save(schools_d, abbrevs)
                                    st.session_state.pop(f"confirm_del_dept_{safe_key}", None)
                                    st.rerun()
                            with nn:
                                if st.button("Cancel", key=f"no_del_dept_{safe_key}"):
                                    st.session_state.pop(f"confirm_del_dept_{safe_key}", None)
                                    st.rerun()

                        # Level editor
                        if st.session_state.get(f"editing_{safe_key}"):
                            with st.form(f"edit_levels_form_{safe_key}"):
                                new_levels = st.number_input(
                                    f"Number of levels for {dept_name}",
                                    min_value=1, max_value=8,
                                    value=int(num_levels), step=1,
                                )
                                sl_save = st.form_submit_button("💾 Save Levels", type="primary")
                                sl_cancel = st.form_submit_button("Cancel")
                            if sl_save:
                                schools_d[sel_school][dept_name] = int(new_levels)
                                if _save(schools_d, abbrevs):
                                    st.success(f"Levels updated to {new_levels}.")
                                    st.session_state.pop(f"editing_{safe_key}", None)
                                    st.rerun()
                                else:
                                    st.error("GitHub write failed.")
                            if sl_cancel:
                                st.session_state.pop(f"editing_{safe_key}", None)
                                st.rerun()

        # ── ICT TAB 4 — Settings ──────────────────────────────────────────────────
        with ict_tab4:
            import datetime as _idt
            st.markdown("### System Settings")
            settings = load_settings()

            # ── School Days & Hours ──────────────────────────────────────────────
            st.markdown("#### 📅 School Days & Hours (WAT)")
            _day_map = {1:"Monday",2:"Tuesday",3:"Wednesday",4:"Thursday",
                        5:"Friday",6:"Saturday",7:"Sunday"}
            _curr_days = settings.get("school_days", [1,2,3,4,5])
            with st.form("ict_school_days"):
                _sel_days = st.multiselect(
                    "School Days",
                    options=list(_day_map.keys()),
                    format_func=lambda d: _day_map[d],
                    default=_curr_days,
                )
                _dc1, _dc2 = st.columns(2)
                with _dc1:
                    _start_t = st.time_input(
                        "School Start (WAT)",
                        value=_idt.time(*map(int, settings.get("school_start","08:30").split(":"))),
                    )
                with _dc2:
                    _end_t = st.time_input(
                        "School End (WAT)",
                        value=_idt.time(*map(int, settings.get("school_end","18:30").split(":"))),
                    )
                _day_btn = st.form_submit_button("💾 Save School Hours", type="primary")
            if _day_btn:
                if not _sel_days:
                    st.error("Select at least one school day.")
                elif _start_t >= _end_t:
                    st.error("Start time must be before end time.")
                else:
                    settings["school_days"]  = sorted(_sel_days)
                    settings["school_start"] = _start_t.strftime("%H:%M")
                    settings["school_end"]   = _end_t.strftime("%H:%M")
                    if save_settings(settings):
                        invalidate_cache("__settings")
                        st.success("School hours saved.")
                    else:
                        st.error("GitHub write failed.")

            st.divider()

            # ── Token Lifetime ────────────────────────────────────────────────────
            st.markdown("#### 🔑 Token Lifetime")
            with st.form("ict_token"):
                tl = st.number_input(
                    "TOKEN_LIFETIME (seconds)",
                    min_value=3, max_value=300, step=1,
                    value=int(settings.get("TOKEN_LIFETIME", 7)),
                    help="How often the 4-digit attendance code rotates",
                )
                st.form_submit_button("💾 Save", type="primary", key="tok_save")
            if st.session_state.get("tok_save"):
                settings["TOKEN_LIFETIME"] = int(tl)
                if save_settings(settings):
                    invalidate_cache("__settings")
                    st.success(f"Saved. TOKEN_LIFETIME = {tl}s")
                else:
                    st.error("GitHub write failed.")

            st.divider()

            # ── Attendance Lifetime ────────────────────────────────────────────────
            st.markdown("#### ⏱ Attendance Lifetime")
            with st.form("ict_att_lifetime"):
                _al1, _al2 = st.columns(2)
                with _al1:
                    st.markdown("**📖 Lecture**")
                    _lec_lt  = st.number_input("Max Duration (minutes)", min_value=5,
                                                max_value=480, step=5,
                                                value=int(settings.get("lecture_lifetime",60)),
                                                key="lec_lt")
                    _lec_act = st.radio("When time expires", ["flag","kill"],
                                        index=0 if settings.get("lecture_action","flag")=="flag" else 1,
                                        format_func=lambda x: "⏰ Flag late entries" if x=="flag" else "⚡ Kill & auto-submit",
                                        key="lec_act")
                with _al2:
                    st.markdown("**🔬 Practical**")
                    _prac_lt  = st.number_input("Max Duration (minutes)", min_value=5,
                                                 max_value=480, step=5,
                                                 value=int(settings.get("practical_lifetime",120)),
                                                 key="prac_lt")
                    _prac_act = st.radio("When time expires", ["flag","kill"],
                                         index=0 if settings.get("practical_action","flag")=="flag" else 1,
                                         format_func=lambda x: "⏰ Flag late entries" if x=="flag" else "⚡ Kill & auto-submit",
                                         key="prac_act")
                _alt_btn = st.form_submit_button("💾 Save Lifetime Settings", type="primary")
            if _alt_btn:
                settings["lecture_lifetime"]   = int(_lec_lt)
                settings["lecture_action"]     = _lec_act
                settings["practical_lifetime"] = int(_prac_lt)
                settings["practical_action"]   = _prac_act
                if save_settings(settings):
                    invalidate_cache("__settings")
                    st.success("Attendance lifetime settings saved.")
                else:
                    st.error("GitHub write failed.")

            st.divider()
            st.markdown(f"""
    | Setting | Value |
    |---------|-------|
    | TOKEN_LIFETIME | **{settings.get('TOKEN_LIFETIME',7)} seconds** |
    | Data Repository | **successcugo/ULASDATA** |
    | LAVA Repository | **successcugo/LAVA** |
    """)
            st.caption("GitHub credentials are in Streamlit Cloud secrets.")

        # ── ICT TAB 5 — Semester Management ──────────────────────────────────────
        with ict_tab5:
            st.markdown("### 📅 Semester Management")
            _active_sem = load_active_semester()

            if _active_sem:
                st.markdown(f"""<div class="info-card">
                    <b>Active Semester:</b> {_active_sem['label']}<br>
                    <b>Started:</b> {_active_sem['started_at'][:16].replace('T',' ')}&nbsp;·&nbsp;
                    <b>By:</b> {_active_sem['started_by']}
                </div>""", unsafe_allow_html=True)
                st.warning("⚠️ Ending the semester will immediately block all student sign-ins and rep session starts.")
                with st.form("end_sem_form"):
                    st.markdown(f"**You are about to end: {_active_sem['label']}**")
                    _end_confirm = st.checkbox("I confirm I want to end this semester")
                    _end_btn = st.form_submit_button("⏹ End Semester", type="primary")
                if _end_btn:
                    if not _end_confirm:
                        st.error("Please tick the confirmation checkbox.")
                    else:
                        _ok, _msg = end_semester(ended_by="ict_master")
                        if _ok:
                            st.success(_msg)
                            st.rerun()
                        else:
                            st.error(_msg)
            else:
                st.info("No semester is currently active. Start one below.")
                with st.form("start_sem_form"):
                    _sem_name = st.selectbox(
                        "Semester",
                        ["First Semester", "Second Semester"],
                        key="new_sem_name",
                    )
                    _sem_session = st.text_input(
                        "Academic Session (e.g. 2025/2026)",
                        placeholder="2025/2026",
                        key="new_sem_session",
                    )
                    _start_btn = st.form_submit_button("▶️ Start Semester", type="primary")
                if _start_btn:
                    import re as _re
                    if not _sem_session.strip():
                        st.error("Please enter the academic session.")
                    elif not _re.match(r"^\d{4}/\d{4}$", _sem_session.strip()):
                        st.error("Session must be in the format YYYY/YYYY e.g. 2025/2026.")
                    else:
                        _ok, _msg = start_semester(_sem_name, _sem_session.strip(), "ict_master")
                        if _ok:
                            if "warning" in _msg.lower():
                                st.warning(_msg)
                            else:
                                st.success(_msg)
                            st.rerun()
                        else:
                            st.error(_msg)

            st.divider()
            st.markdown("#### 📜 Semester History")
            _sem_hist = load_semester_history()
            if not _sem_hist:
                st.caption("No completed semesters yet.")
            else:
                import pandas as _spd
                _sh_rows = []
                for _s in _sem_hist:
                    _sh_rows.append({
                        "Semester":   _s.get("label", ""),
                        "Started":    _s.get("started_at", "")[:16].replace("T", " "),
                        "Ended":      _s.get("ended_at",   "")[:16].replace("T", " "),
                        "Started By": _s.get("started_by", ""),
                        "Ended By":   _s.get("ended_by",   ""),
                    })
                st.dataframe(_spd.DataFrame(_sh_rows), use_container_width=True, hide_index=True)

        st.stop()


    # ═══════════════════════════════════════════════════════════════════════════════
    #  ADVISOR DASHBOARD
    # ═══════════════════════════════════════════════════════════════════════════════
    if st.session_state.portal_role == "advisor":

        adv        = st.session_state.adv_user
        school     = adv["school"]
        department = adv["department"]

        hc, lc = st.columns([5, 1])
        with hc:
            st.markdown(f"""<div class="info-card">
                <b>Advisor:</b> {adv['username']} &nbsp;·&nbsp;
                <b>Dept:</b> {department} &nbsp;·&nbsp;
                <b>School:</b> {school[:35]}
            </div>""", unsafe_allow_html=True)
        with lc:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Logout", key="adv_logout"):
                do_logout()
                st.rerun()

        tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["👥 Course Reps", "🔑 Passwords", "🏷️ Abbreviation", "📊 GPA Calculator", "🎓 GPA Management", "📈 Dept Stats"])

        # ── TAB 1 — Course Reps ───────────────────────────────────────────────────
        with tab1:
            st.markdown(f"### Course Reps — {department}")
            invalidate_cache("__users")
            reps     = get_reps_for_dept(school, department)
            levels   = get_levels(department, school)

            if reps:
                # Group reps by level for display
                from collections import defaultdict as _dd
                reps_by_level = _dd(list)
                for rep in reps:
                    reps_by_level[rep["level"]].append(rep)

                for lvl in sorted(reps_by_level.keys()):
                    st.markdown(f"**Level {lvl}L**")
                    for rep in reps_by_level[lvl]:
                        rc, dc = st.columns([6, 1])
                        with rc:
                            st.markdown(f"""
                            <div class="rep-card">
                                <span class="badge-rep">Level {rep['level']}L</span>
                                &nbsp;&nbsp;<b>{rep['username']}</b>
                                &nbsp;·&nbsp;
                                <span style="font-size:0.85rem;opacity:0.65">
                                    Created {rep.get('created_at','')[:10]}
                                    by {rep.get('created_by','—')}
                                </span>
                            </div>""", unsafe_allow_html=True)
                        with dc:
                            if st.button("🗑️", key=f"del_rep_{rep['username']}", help=f"Remove {rep['username']}"):
                                delete_user(rep["username"])
                                invalidate_cache("__users")
                                st.success(f"Removed {rep['username']}")
                                st.rerun()
            else:
                st.info("No course reps assigned yet for this department.")

            st.divider()
            st.markdown("### ➕ Add Course Rep")
            st.caption("Multiple reps can share a level — each gets their own login.")
            with st.form("assign_rep"):
                level = st.selectbox("Level", levels)
                nu    = st.text_input("Username (must be unique across all of FUTO)")
                np    = st.text_input("Password", type="password")
                np2   = st.text_input("Confirm Password", type="password")
                a_btn = st.form_submit_button("Add Rep", type="primary")
            if a_btn:
                if not nu.strip() or not np.strip():
                    st.error("All fields are required.")
                elif np != np2:
                    st.error("Passwords do not match.")
                elif len(np) < 6:
                    st.error("Password must be at least 6 characters.")
                else:
                    ok, msg = create_user(nu.strip(), np, "rep", school, department, level, adv["username"])
                    if ok:
                        invalidate_cache("__users")
                        st.success(f"✅ Rep '{nu}' added to Level {level}L!")
                        st.rerun()
                    else:
                        st.error(msg)

        # ── TAB 2 — Passwords ─────────────────────────────────────────────────────
        with tab2:
            invalidate_cache("__users")
            reps     = get_reps_for_dept(school, department)
            co_advs  = get_advisors_for_dept(school, department)
            # Co-advisors = all advisors in same dept except self
            co_advs  = [a for a in co_advs if a["username"] != adv["username"]]

            # ── Reset a rep's password ────────────────────────────────────────────
            st.markdown("### Change a Course Rep's Password")
            if not reps:
                st.info("No course reps in this department.")
            else:
                rep_opts  = {f"{r['username']} (Level {r['level']}L)": r["username"]
                             for r in sorted(reps, key=lambda r: r["level"])}
                sel_lbl   = st.selectbox("Select Rep", list(rep_opts.keys()), key="pw_sel")
                sel_uname = rep_opts[sel_lbl]
                with st.form("rep_pw"):
                    new_p  = st.text_input("New Password", type="password")
                    new_p2 = st.text_input("Confirm New Password", type="password")
                    pw_btn = st.form_submit_button("Update Password", type="primary")
                if pw_btn:
                    if not new_p.strip():
                        st.error("Password cannot be empty.")
                    elif new_p != new_p2:
                        st.error("Passwords do not match.")
                    elif len(new_p) < 6:
                        st.error("Password must be at least 6 characters.")
                    else:
                        update_password(sel_uname, new_p)
                        invalidate_cache("__users")
                        st.success(f"Password updated for {sel_uname}.")

            # ── Reset a co-advisor's password ─────────────────────────────────────
            st.divider()
            st.markdown("### Change a Co-Advisor's Password")
            st.caption("You can reset passwords for other advisors in your department.")
            if not co_advs:
                st.info("No other advisors in your department.")
            else:
                co_opts   = {a["username"]: a["username"] for a in co_advs}
                co_sel    = st.selectbox("Select Co-Advisor", list(co_opts.keys()), key="co_pw_sel")
                with st.form("co_adv_pw"):
                    co_new  = st.text_input("New Password", type="password")
                    co_new2 = st.text_input("Confirm New Password", type="password")
                    co_btn  = st.form_submit_button("Update Co-Advisor Password", type="primary")
                if co_btn:
                    if not co_new.strip():
                        st.error("Password cannot be empty.")
                    elif co_new != co_new2:
                        st.error("Passwords do not match.")
                    elif len(co_new) < 6:
                        st.error("Password must be at least 6 characters.")
                    else:
                        update_password(co_sel, co_new)
                        invalidate_cache("__users")
                        st.success(f"Password updated for {co_sel}.")

            # ── Change own password ────────────────────────────────────────────────
            st.divider()
            st.markdown("### Change My Password")
            with st.form("adv_pw"):
                old_p  = st.text_input("Current Password", type="password")
                my_new = st.text_input("New Password", type="password")
                my_n2  = st.text_input("Confirm New Password", type="password")
                ap_btn = st.form_submit_button("Change My Password", type="primary")
            if ap_btn:
                if not verify_password(old_p, adv["password_hash"]):
                    st.error("Current password is incorrect.")
                elif my_new != my_n2:
                    st.error("Passwords do not match.")
                elif len(my_new) < 6:
                    st.error("Password must be at least 6 characters.")
                else:
                    update_password(adv["username"], my_new)
                    st.session_state.adv_user["password_hash"] = hash_password(my_new)
                    invalidate_cache("__users")
                    st.success("Password updated.")



            st.divider()
            # ── Attendance Lifetime Overrides (per level) ─────────────────────
            st.markdown("### ⏱ Attendance Lifetime Settings")
            st.caption("Set custom attendance durations for each level in your department. Cannot exceed ICT maximums.")

            _ict_settings = load_settings()
            _ict_lec_max  = int(_ict_settings.get("lecture_lifetime",  60))
            _ict_prac_max = int(_ict_settings.get("practical_lifetime", 120))

            _adv_levels = sorted(get_levels(department, school))
            if not _adv_levels:
                st.info("No levels configured for your department yet.")
            else:
                _sel_adv_lvl = st.selectbox("Level", _adv_levels,
                                             format_func=lambda l: f"{l}L",
                                             key="adv_lt_level")
                _curr_lt = load_advisor_lifetime(school, department, _sel_adv_lvl)

                with st.form("adv_lt_form"):
                    _alt1, _alt2 = st.columns(2)
                    with _alt1:
                        st.markdown(f"**📖 Lecture** (ICT max: {_ict_lec_max} min)")
                        _adv_lec = st.number_input(
                            "Duration (minutes)",
                            min_value=5, max_value=_ict_lec_max, step=5,
                            value=int(_curr_lt.get("lecture_lifetime", _ict_lec_max)),
                            key="adv_lec_lt",
                        )
                    with _alt2:
                        st.markdown(f"**🔬 Practical** (ICT max: {_ict_prac_max} min)")
                        _adv_prac = st.number_input(
                            "Duration (minutes)",
                            min_value=5, max_value=_ict_prac_max, step=5,
                            value=int(_curr_lt.get("practical_lifetime", _ict_prac_max)),
                            key="adv_prac_lt",
                        )
                    _adv_lt_btn = st.form_submit_button("💾 Save", type="primary")

                if _adv_lt_btn:
                    _new_lt = {
                        "lecture_lifetime":   min(int(_adv_lec),  _ict_lec_max),
                        "practical_lifetime": min(int(_adv_prac), _ict_prac_max),
                        "set_by": adv["username"],
                    }
                    if save_advisor_lifetime(school, department, _sel_adv_lvl, _new_lt):
                        st.success(f"Saved for Level {_sel_adv_lvl}L — "
                                   f"Lecture: {_new_lt['lecture_lifetime']}min, "
                                   f"Practical: {_new_lt['practical_lifetime']}min")
                    else:
                        st.error("GitHub write failed.")

        # ── TAB 3 — Department Abbreviation ──────────────────────────────────────
        with tab3:
            st.markdown(f"### Department Abbreviation — {department}")
            st.markdown(
                "Set a **3-letter abbreviation** for your department. "
                "This is used in attendance file names pushed to LAVA, "
                "e.g. `SEETCSC**EEE**300_2026-01-15.csv`."
            )

            current_abbr = get_dept_abbreviation(department)
            st.markdown(f"""<div class="info-card">
                <b>Current abbreviation:</b> &nbsp;
                <span style="font-size:1.4rem;font-weight:900;letter-spacing:2px">{current_abbr}</span>
                &nbsp;&nbsp;<span style="opacity:0.6;font-size:0.85rem">(used in all future file names)</span>
            </div>""", unsafe_allow_html=True)

            with st.form("abbr_form"):
                new_abbr = st.text_input(
                    "Abbreviation (any length, letters only)",
                    value=current_abbr,
                    placeholder="e.g. EEE, EPSENG, POWER",
                )
                abbr_btn = st.form_submit_button("💾 Save Abbreviation", type="primary")

            if abbr_btn:
                cleaned = new_abbr.strip().upper()
                if not cleaned:
                    st.error("Abbreviation cannot be empty.")
                elif not cleaned.isalpha():
                    st.error("Abbreviation must contain letters only (no spaces or symbols).")
                else:
                    ok = set_dept_abbreviation(department, cleaned)
                    if ok:
                        invalidate_cache("__settings")
                        st.success(f"✅ Abbreviation for **{department}** set to **{cleaned}**.")
                        st.caption("All attendance files pushed from now on will use this abbreviation.")
                    else:
                        st.error("Failed to save — GitHub write error.")

            st.divider()
            st.markdown("#### Example filename with this abbreviation")
            example = f"SEET · CSC301 · {current_abbr} · Level 300 · 2026-01-15"
            st.markdown(f"`SEETCSC301{current_abbr}300_2026-01-15.csv`")
            st.caption(f"Format: SCHOOL + COURSECODE + ABBR + LEVEL + _ + DATE .csv")

        # ── TAB 4 — GPA Calculator ───────────────────────────────────────────────
        with tab4:
            st.markdown("### 📊 GPA Calculator")
            st.markdown(
                "Enter each course with its unit load and grade. "
                "GPA = Total Weighted Points ÷ Total Units."
            )

            # FUTO grade → point mapping
            GRADE_POINTS = {"A": 5.0, "B": 4.0, "C": 3.0, "D": 2.0, "E": 1.0, "F": 0.0}

            # Grading reference
            with st.expander("📋 FUTO Grading Scale"):
                st.markdown("""
    | Grade | Point |
    |-------|-------|
    | A | 5.0 |
    | B | 4.0 |
    | C | 3.0 |
    | D | 2.0 |
    | E | 1.0 |
    | F | 0.0 |
    """)

            # Keep courses in session state so they persist across reruns
            if "gpa_courses" not in st.session_state:
                st.session_state.gpa_courses = [
                    {"code": "", "units": 1, "grade": "A"}
                ]

            st.markdown("#### Course Entries")

            # Render each course row
            to_delete = None
            for i, course in enumerate(st.session_state.gpa_courses):
                c1, c2, c3, c4 = st.columns([3, 1, 2, 1])
                with c1:
                    st.session_state.gpa_courses[i]["code"] = st.text_input(
                        "Course Code", value=course["code"],
                        key=f"gpa_code_{i}", placeholder="e.g. CSC301",
                        label_visibility="collapsed" if i > 0 else "visible",
                    )
                with c2:
                    st.session_state.gpa_courses[i]["units"] = st.number_input(
                        "Units", min_value=1, max_value=6, step=1,
                        value=course["units"], key=f"gpa_units_{i}",
                        label_visibility="collapsed" if i > 0 else "visible",
                    )
                with c3:
                    st.session_state.gpa_courses[i]["grade"] = st.selectbox(
                        "Grade", list(GRADE_POINTS.keys()),
                        index=list(GRADE_POINTS.keys()).index(course["grade"]),
                        key=f"gpa_grade_{i}",
                        label_visibility="collapsed" if i > 0 else "visible",
                    )
                with c4:
                    st.markdown("<br>" if i == 0 else "", unsafe_allow_html=True)
                    if len(st.session_state.gpa_courses) > 1:
                        if st.button("✕", key=f"gpa_del_{i}", help="Remove this course"):
                            to_delete = i

            if to_delete is not None:
                st.session_state.gpa_courses.pop(to_delete)
                st.rerun()

            col_add, col_clear = st.columns([1, 1])
            with col_add:
                if st.button("➕ Add Course", use_container_width=True):
                    st.session_state.gpa_courses.append({"code": "", "units": 1, "grade": "A"})
                    st.rerun()
            with col_clear:
                if st.button("🗑️ Clear All", use_container_width=True):
                    st.session_state.gpa_courses = [{"code": "", "units": 1, "grade": "A"}]
                    st.rerun()

            st.divider()

            # ── Calculation ───────────────────────────────────────────────────────
            courses = st.session_state.gpa_courses
            valid   = [c for c in courses if str(c["code"]).strip()]

            if not valid:
                st.info("Add at least one course with a course code to calculate GPA.")
            else:
                total_units    = sum(c["units"] for c in valid)
                total_weighted = sum(c["units"] * GRADE_POINTS[c["grade"]] for c in valid)
                gpa            = total_weighted / total_units if total_units > 0 else 0.0

                # Colour based on GPA range
                if gpa >= 4.5:
                    gpa_colour, standing = "#27ae60", "First Class"
                elif gpa >= 3.5:
                    gpa_colour, standing = "#2980b9", "Second Class Upper"
                elif gpa >= 2.4:
                    gpa_colour, standing = "#f39c12", "Second Class Lower"
                elif gpa >= 1.5:
                    gpa_colour, standing = "#e67e22", "Third Class"
                else:
                    gpa_colour, standing = "#c0392b", "Pass / Fail"

                st.markdown(f"""
                <div style="
                    background: rgba(0,0,0,0.05);
                    border: 2px solid {gpa_colour};
                    border-radius: 14px;
                    padding: 1.5rem;
                    text-align: center;
                    margin: 1rem 0;
                    color: inherit;
                ">
                    <div style="font-size:3rem;font-weight:900;color:{gpa_colour};line-height:1">
                        {gpa:.2f}
                    </div>
                    <div style="font-size:1rem;font-weight:700;color:{gpa_colour};margin-top:0.3rem">
                        {standing}
                    </div>
                    <div style="font-size:0.8rem;opacity:0.65;margin-top:0.5rem">
                        {total_weighted:.1f} weighted points ÷ {total_units} units
                    </div>
                </div>
                """, unsafe_allow_html=True)

                # Breakdown table
                st.markdown("#### Breakdown")
                import pandas as pd
                rows = []
                for c in valid:
                    wp = c["units"] * GRADE_POINTS[c["grade"]]
                    rows.append({
                        "Course Code":     c["code"].upper(),
                        "Units":           c["units"],
                        "Grade":           c["grade"],
                        "Grade Point":     GRADE_POINTS[c["grade"]],
                        "Weighted Points": f"{GRADE_POINTS[c['grade']]} × {c['units']} = {wp:.0f}",
                    })
                st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

                wp_parts = " + ".join(
                    str(int(c["units"] * GRADE_POINTS[c["grade"]])) for c in valid
                )
                st.markdown(
                    f"**Total Weighted Points** = {wp_parts} = **{total_weighted:.0f}**\n\n"
                    f"**Total Units** = {total_units}\n\n"
                    f"**GPA** = {total_weighted:.0f} ÷ {total_units} = **{gpa:.2f}**"
                )


                # ── Assign GPA to Student ─────────────────────────────────────────────
                st.divider()
                st.markdown("### 🎓 Assign Semester GPA to Student")

                _active_sem_gpa = load_active_semester()
                if not _active_sem_gpa:
                    st.warning("⚠️ No active semester. Ask ICT to start a semester before assigning GPAs.")
                else:
                    st.markdown(f"""<div class="info-card">
                        <b>Assigning for:</b> {_active_sem_gpa['label']}
                    </div>""", unsafe_allow_html=True)

                    # Load matric list lazily — cached in session state
                    if "gpa_matric_list" not in st.session_state:
                        with st.spinner("Loading student matric numbers from LAVA..."):
                            try:
                                st.session_state.gpa_matric_list = get_dept_matric_numbers(school, department)
                            except Exception:
                                st.session_state.gpa_matric_list = []

                    _matric_list = st.session_state.gpa_matric_list

                    if not _matric_list:
                        st.info("No student matric numbers found in LAVA records for your department yet.")
                    else:
                        _ga1, _ga2 = st.columns([3, 1])
                        with _ga1:
                            _sel_matric = st.selectbox(
                                "Select Student Matric Number",
                                _matric_list,
                                key="gpa_assign_matric",
                                help="Only students from your department's LAVA records are shown.",
                            )
                        with _ga2:
                            if st.button("🔄 Refresh List", key="gpa_refresh_matric",
                                         use_container_width=True):
                                st.session_state.pop("gpa_matric_list", None)
                                st.rerun()

                        # Show existing CGPA for selected student
                        _existing_records = load_student_gpa(_sel_matric)
                        if _existing_records:
                            _cgpa = compute_cgpa(_existing_records)
                            _cgpa_colour = (
                                "#27ae60" if _cgpa >= 4.5 else
                                "#2980b9" if _cgpa >= 3.5 else
                                "#f39c12" if _cgpa >= 2.4 else
                                "#e67e22" if _cgpa >= 1.5 else "#c0392b"
                            )
                            _cgpa_standing = (
                                "First Class" if _cgpa >= 4.5 else
                                "Second Class Upper" if _cgpa >= 3.5 else
                                "Second Class Lower" if _cgpa >= 2.4 else
                                "Third Class" if _cgpa >= 1.5 else "Pass / Fail"
                            )
                            st.markdown(f"""
                            <div style="display:flex;gap:1rem;margin:0.8rem 0;flex-wrap:wrap">
                                <div style="background:rgba(0,0,0,0.05);border:2px solid {_cgpa_colour};
                                    border-radius:12px;padding:1rem 1.5rem;text-align:center;min-width:140px">
                                    <div style="font-size:2.2rem;font-weight:900;color:{_cgpa_colour};line-height:1">
                                        {_cgpa:.2f}
                                    </div>
                                    <div style="font-size:0.75rem;opacity:0.7;text-transform:uppercase;
                                        letter-spacing:1px;margin-top:0.3rem">CGPA</div>
                                    <div style="font-size:0.8rem;color:{_cgpa_colour};margin-top:0.2rem;
                                        font-weight:600">{_cgpa_standing}</div>
                                </div>
                                <div style="display:flex;align-items:center;font-size:0.9rem;opacity:0.7">
                                    Based on {len(_existing_records)} semester(s)
                                </div>
                            </div>
                            """, unsafe_allow_html=True)

                            # Collapsible history table
                            with st.expander(f"📋 Semester history for {_sel_matric}"):
                                import pandas as _gpd
                                _hist_rows = []
                                for _r in _existing_records:
                                    _hist_rows.append({
                                        "Semester":    _r.get("semester", ""),
                                        "GPA":         _r.get("gpa", ""),
                                        "Assigned By": _r.get("assigned_by", ""),
                                        "Date":        _r.get("assigned_at", "")[:10],
                                    })
                                st.dataframe(
                                    _gpd.DataFrame(_hist_rows),
                                    use_container_width=True, hide_index=True
                                )
                        else:
                            st.caption(f"No GPA records yet for matric {_sel_matric}.")

                        # Check if this semester already assigned
                        _already = next(
                            (r for r in _existing_records if r.get("semester") == _active_sem_gpa["label"]),
                            None
                        )
                        if _already:
                            st.info(
                                f"A GPA of **{_already['gpa']}** has already been assigned for "
                                f"**{_active_sem_gpa['label']}**. Submit again to update it."
                            )

                        # Assignment form
                        with st.form("gpa_assign_form"):
                            _new_gpa = st.number_input(
                                f"GPA for {_active_sem_gpa['label']}",
                                min_value=0.0, max_value=5.0,
                                step=0.01, format="%.2f",
                                value=float(_already["gpa"]) if _already else 0.0,
                                key="gpa_assign_value",
                            )
                            _assign_btn = st.form_submit_button(
                                "💾 Save GPA", type="primary", use_container_width=True
                            )
                        if _assign_btn:
                            _ok_g, _msg_g = assign_semester_gpa(
                                _sel_matric, _new_gpa,
                                adv["username"], department
                            )
                            if _ok_g:
                                st.success(_msg_g)
                                # Refresh matric list cache for updated CGPA display
                                st.session_state.pop("gpa_matric_list", None)
                                st.rerun()
                            else:
                                st.error(_msg_g)

        # ── TAB 5 — GPA Management ───────────────────────────────────────────────
        with tab5:
            import pandas as _gpd
            from io import BytesIO as _BytesIO

            st.markdown(f"### 🎓 GPA Management — {department}")

            # ── Load student roster from LAVA (cached per session) ───────────────
            if "gpa_mgmt_students" not in st.session_state:
                with st.spinner("Loading student records from LAVA…"):
                    try:
                        st.session_state.gpa_mgmt_students = get_dept_students(school, department)
                    except Exception:
                        st.session_state.gpa_mgmt_students = []

            _all_students = st.session_state.gpa_mgmt_students

            # ── Filters row ───────────────────────────────────────────────────────
            _gf1, _gf2, _gf3 = st.columns([3, 2, 1])
            with _gf1:
                _search = st.text_input("🔍 Search name or matric", key="gpa_search",
                                        placeholder="e.g. Adaeze or 2021/12345")
            with _gf2:
                _levels = sorted({s["level"] for s in _all_students}) if _all_students else []
                _lvl_opts = ["All Levels"] + [f"{l}L" for l in _levels]
                _sel_lvl = st.selectbox("Filter by Level", _lvl_opts, key="gpa_lvl_filter")
            with _gf3:
                st.markdown("<div style='padding-top:1.75rem'>", unsafe_allow_html=True)
                if st.button("🔄 Refresh", key="gpa_refresh_roster", use_container_width=True):
                    st.session_state.pop("gpa_mgmt_students", None)
                    st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)

            # ── Apply filters ─────────────────────────────────────────────────────
            _filtered = _all_students
            if _sel_lvl != "All Levels":
                _lvl_code = _sel_lvl.replace("L", "")
                _filtered = [s for s in _filtered if s["level"] == _lvl_code]
            if _search.strip():
                _q = _search.strip().lower()
                _filtered = [s for s in _filtered
                             if _q in s["full_name"].lower() or _q in s["matric"].lower()]

            if not _all_students:
                st.info("No student records found in LAVA for your department yet. "
                        "Push some attendance first.")
                st.stop()

            if not _filtered:
                st.warning("No students match your search.")
                st.stop()

            # ── Student selector ──────────────────────────────────────────────────
            _stu_labels = [f"{s['full_name']} — {s['matric']} ({s['level']}L)"
                           for s in _filtered]
            _sel_idx = st.selectbox("Select Student", range(len(_stu_labels)),
                                    format_func=lambda i: _stu_labels[i],
                                    key="gpa_sel_student")
            _sel_stu = _filtered[_sel_idx]
            _sel_matric = _sel_stu["matric"]

            st.markdown("---")

            # ── GPA history for selected student ──────────────────────────────────
            _gpa_records = load_student_gpa(_sel_matric)

            _cgpa_col, _hist_col = st.columns([1, 2])
            with _cgpa_col:
                if _gpa_records:
                    _cgpa = compute_cgpa(_gpa_records)
                    _standing = (
                        "First Class"        if _cgpa >= 4.5 else
                        "Second Class Upper" if _cgpa >= 3.5 else
                        "Second Class Lower" if _cgpa >= 2.4 else
                        "Third Class"        if _cgpa >= 1.5 else "Pass / Fail"
                    )
                    _col = (
                        "#27ae60" if _cgpa >= 4.5 else
                        "#2980b9" if _cgpa >= 3.5 else
                        "#f39c12" if _cgpa >= 2.4 else
                        "#e67e22" if _cgpa >= 1.5 else "#c0392b"
                    )
                    st.markdown(f"""
                    <div style="background:rgba(0,0,0,0.05);border:2px solid {_col};
                        border-radius:14px;padding:1.2rem;text-align:center">
                        <div style="font-size:2.8rem;font-weight:900;color:{_col};line-height:1">
                            {_cgpa:.2f}
                        </div>
                        <div style="font-size:0.7rem;opacity:0.6;text-transform:uppercase;
                            letter-spacing:1px;margin-top:0.3rem">CGPA</div>
                        <div style="font-size:0.85rem;color:{_col};font-weight:600;
                            margin-top:0.25rem">{_standing}</div>
                        <div style="font-size:0.75rem;opacity:0.55;margin-top:0.2rem">
                            {len(_gpa_records)} semester(s)
                        </div>
                    </div>""", unsafe_allow_html=True)
                else:
                    st.markdown("""
                    <div style="background:rgba(0,0,0,0.05);border:2px dashed rgba(128,128,128,0.3);
                        border-radius:14px;padding:1.2rem;text-align:center;opacity:0.6">
                        <div style="font-size:2rem">—</div>
                        <div style="font-size:0.8rem">No GPA recorded yet</div>
                    </div>""", unsafe_allow_html=True)

            with _hist_col:
                if _gpa_records:
                    _hist_df = _gpd.DataFrame([{
                        "Semester":    r.get("semester", ""),
                        "GPA":         f"{r.get('gpa', 0):.2f}",
                        "Assigned By": r.get("assigned_by", ""),
                        "Date":        r.get("assigned_at", "")[:10],
                    } for r in _gpa_records])
                    st.dataframe(_hist_df, use_container_width=True, hide_index=True)
                else:
                    st.caption("GPA history will appear here once recorded.")

            # ── Assign GPA form ───────────────────────────────────────────────────
            st.markdown("#### Assign / Update GPA")

            _sessions = get_available_sessions()
            if not _sessions:
                st.warning("No sessions found. ICT must start at least one semester first.")
            else:
                _ga1, _ga2 = st.columns(2)
                with _ga1:
                    _sel_session = st.selectbox("Academic Session", _sessions,
                                                key="gpa_assign_session")
                with _ga2:
                    _session_sems = get_semesters_for_session(_sel_session)
                    if not _session_sems:
                        st.warning(f"No semesters found for {_sel_session}.")
                        _sel_sem_record = None
                    else:
                        _sem_labels = [s["label"] for s in _session_sems]
                        _sel_sem_label = st.selectbox("Semester", _sem_labels,
                                                      key="gpa_assign_semester")
                        _sel_sem_record = next(
                            (s for s in _session_sems if s["label"] == _sel_sem_label), None
                        )

                if _sel_sem_record:
                    # Check if GPA already assigned for this semester
                    _existing = next(
                        (r for r in _gpa_records
                         if r.get("semester") == _sel_sem_record["label"]),
                        None
                    )
                    if _existing:
                        st.info(f"GPA of **{_existing['gpa']:.2f}** already assigned for "
                                f"**{_sel_sem_record['label']}**. Submit to update.")

                    with st.form("gpa_mgmt_assign_form"):
                        _new_gpa = st.number_input(
                            f"GPA for {_sel_sem_record['label']}",
                            min_value=0.0, max_value=5.0, step=0.01, format="%.2f",
                            value=float(_existing["gpa"]) if _existing else 0.0,
                        )
                        _save_btn = st.form_submit_button(
                            "💾 Save GPA", type="primary", use_container_width=True
                        )

                    if _save_btn:
                        _ok, _msg = assign_gpa_for_semester(
                            _sel_matric, _new_gpa,
                            adv["username"], department,
                            _sel_sem_record["label"],
                        )
                        if _ok:
                            st.success(_msg)
                            st.rerun()
                        else:
                            st.error(_msg)

            # ══════════════════════════════════════════════════════════════════════
            # ── Level Export ─────────────────────────────────────────────────────
            # ══════════════════════════════════════════════════════════════════════
            st.markdown("---")
            st.markdown("### 📤 Export Level GPA List")

            _exp_levels = sorted({s["level"] for s in _all_students})
            if not _exp_levels:
                st.info("No student records available for export.")
            else:
                _exp_lvl = st.selectbox(
                    "Select Level to Export",
                    [f"{l}L" for l in _exp_levels],
                    key="gpa_export_level"
                )
                _exp_lvl_code = _exp_lvl.replace("L", "")
                _exp_students = [s for s in _all_students if s["level"] == _exp_lvl_code]

                # Build export rows — one row per student with all semester GPAs + CGPA
                def _build_export_rows():
                    # Gather all unique semester labels across all students
                    all_sem_labels: list[str] = []
                    all_gpa_data: dict[str, list] = {}
                    for _s in _exp_students:
                        _recs = load_student_gpa(_s["matric"])
                        all_gpa_data[_s["matric"]] = _recs
                        for _r in _recs:
                            _lbl = _r.get("semester", "")
                            if _lbl and _lbl not in all_sem_labels:
                                all_sem_labels.append(_lbl)

                    # Sort semester labels chronologically by their started_at if possible
                    # Fall back to alphabetical otherwise
                    all_sem_labels = sorted(set(all_sem_labels))

                    rows = []
                    for _s in _exp_students:
                        _recs = all_gpa_data[_s["matric"]]
                        _rec_map = {r["semester"]: r["gpa"] for r in _recs}
                        _cgpa_val = compute_cgpa(_recs) if _recs else None
                        row = {
                            "S/N":        len(rows) + 1,
                            "Surname":    _s["surname"],
                            "Other Names": _s["other_names"],
                            "Matric No.": _s["matric"],
                        }
                        for _lbl in all_sem_labels:
                            row[_lbl] = f"{_rec_map[_lbl]:.2f}" if _lbl in _rec_map else "—"
                        row["CGPA"] = f"{_cgpa_val:.2f}" if _cgpa_val is not None else "—"
                        rows.append(row)
                    return rows, all_sem_labels

                _ecol1, _ecol2, _ecol3 = st.columns(3)

                # ── Excel export ──────────────────────────────────────────────────
                def _export_excel():
                    import openpyxl as _xl
                    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
                    rows, sem_labels = _build_export_rows()
                    wb = _xl.Workbook()
                    ws = wb.active
                    ws.title = f"GPA {_exp_lvl}"

                    RED_FILL  = PatternFill("solid", fgColor="7B0000")
                    GREY_FILL = PatternFill("solid", fgColor="F2F2F2")
                    thin      = Side(style="thin", color="CCCCCC")
                    border    = Border(left=thin, right=thin, top=thin, bottom=thin)

                    # Title
                    ws.merge_cells(f"A1:{chr(68 + len(sem_labels))}1")
                    title_cell = ws["A1"]
                    title_cell.value = (f"Federal University of Technology, Owerri\n"
                                        f"{department} — {_exp_lvl} GPA Record")
                    title_cell.font      = Font(bold=True, size=13, color="7B0000")
                    title_cell.alignment = Alignment(wrap_text=True, horizontal="center",
                                                     vertical="center")
                    ws.row_dimensions[1].height = 42

                    # Headers
                    headers = ["S/N", "Surname", "Other Names", "Matric No."] + sem_labels + ["CGPA"]
                    for ci, h in enumerate(headers, 1):
                        cell = ws.cell(row=2, column=ci, value=h)
                        cell.font      = Font(bold=True, color="FFFFFF", size=10)
                        cell.fill      = RED_FILL
                        cell.alignment = Alignment(horizontal="center", vertical="center",
                                                   wrap_text=True)
                        cell.border    = border
                    ws.row_dimensions[2].height = 32

                    # Data rows
                    for ri, row in enumerate(rows, 3):
                        fill = GREY_FILL if ri % 2 == 0 else PatternFill()
                        for ci, h in enumerate(headers, 1):
                            cell = ws.cell(row=ri, column=ci, value=row.get(h, ""))
                            cell.fill      = fill
                            cell.border    = border
                            cell.alignment = Alignment(horizontal="center" if ci != 3 else "left",
                                                       vertical="center")
                            cell.font      = Font(size=10,
                                                  bold=(h == "CGPA"),
                                                  color=("7B0000" if h == "CGPA" else "000000"))

                    # Column widths
                    col_widths = [5, 18, 18, 16] + [22] * len(sem_labels) + [10]
                    for ci, w in enumerate(col_widths, 1):
                        ws.column_dimensions[_xl.utils.get_column_letter(ci)].width = w

                    buf = _BytesIO()
                    wb.save(buf)
                    return buf.getvalue()

                # ── Word export ───────────────────────────────────────────────────
                def _export_word():
                    from docx import Document as _Doc
                    from docx.shared import Pt, RGBColor, Cm
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    rows, sem_labels = _build_export_rows()
                    doc = _Doc()
                    for sec in doc.sections:
                        sec.page_width  = Cm(29.7)
                        sec.page_height = Cm(21.0)
                        sec.left_margin = sec.right_margin = Cm(1.8)

                    def _add_run(para, text, bold=False, size=11, color=None):
                        run = para.add_run(text)
                        run.bold      = bold
                        run.font.size = Pt(size)
                        if color:
                            run.font.color.rgb = RGBColor(*color)
                        return run

                    # Title block
                    t = doc.add_paragraph()
                    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(t, "Federal University of Technology, Owerri\n",
                             bold=True, size=14, color=(123, 0, 0))
                    _add_run(t, f"{department} — {_exp_lvl} GPA Record",
                             bold=True, size=12, color=(123, 0, 0))

                    doc.add_paragraph()

                    # Table
                    headers = ["S/N", "Surname", "Other Names", "Matric No."] + sem_labels + ["CGPA"]
                    tbl = doc.add_table(rows=1, cols=len(headers))
                    tbl.style = "Table Grid"

                    def _set_cell_bg(cell, hex_color):
                        tc   = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), hex_color)
                        tcPr.append(shd)

                    # Header row
                    hdr_row = tbl.rows[0]
                    for i, h in enumerate(headers):
                        cell = hdr_row.cells[i]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _add_run(p, h, bold=True, size=9, color=(255, 255, 255))
                        _set_cell_bg(cell, "7B0000")

                    # Data rows
                    for ri, row in enumerate(rows):
                        tr = tbl.add_row()
                        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
                        for ci, h in enumerate(headers):
                            cell = tr.cells[ci]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                                           if h in ("Surname", "Other Names")
                                           else WD_ALIGN_PARAGRAPH.CENTER)
                            _add_run(p, str(row.get(h, "")),
                                     bold=(h == "CGPA"), size=9,
                                     color=((123, 0, 0) if h == "CGPA" else (0, 0, 0)))
                            _set_cell_bg(cell, bg)

                    buf = _BytesIO()
                    doc.save(buf)
                    return buf.getvalue()

                # ── PDF export ────────────────────────────────────────────────────
                def _export_pdf():
                    from reportlab.lib.pagesizes import A4, landscape
                    from reportlab.platypus import (SimpleDocTemplate, Table,
                                                    TableStyle, Paragraph, Spacer)
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib import colors
                    from reportlab.lib.units import cm

                    rows, sem_labels = _build_export_rows()
                    buf   = _BytesIO()
                    doc   = SimpleDocTemplate(buf, pagesize=landscape(A4),
                                              leftMargin=1.5*cm, rightMargin=1.5*cm,
                                              topMargin=1.5*cm, bottomMargin=1.5*cm)
                    styles = getSampleStyleSheet()
                    RED    = colors.HexColor("#7B0000")
                    LGREY  = colors.HexColor("#F2F2F2")

                    title_style = ParagraphStyle("title", parent=styles["Normal"],
                                                 fontSize=13, textColor=RED,
                                                 fontName="Helvetica-Bold",
                                                 alignment=1, spaceAfter=6)
                    sub_style   = ParagraphStyle("sub", parent=styles["Normal"],
                                                 fontSize=10, textColor=RED,
                                                 fontName="Helvetica-Bold",
                                                 alignment=1, spaceAfter=12)

                    headers = ["S/N", "Surname", "Other Names", "Matric No."] + sem_labels + ["CGPA"]
                    tbl_data = [headers]
                    for row in rows:
                        tbl_data.append([str(row.get(h, "")) for h in headers])

                    col_w = ([1.0*cm, 3.5*cm, 3.5*cm, 3.0*cm]
                             + [3.2*cm] * len(sem_labels)
                             + [2.0*cm])

                    tbl_obj = Table(tbl_data, colWidths=col_w, repeatRows=1)
                    tbl_style = [
                        ("BACKGROUND",  (0,0), (-1,0), RED),
                        ("TEXTCOLOR",   (0,0), (-1,0), colors.white),
                        ("FONTNAME",    (0,0), (-1,0), "Helvetica-Bold"),
                        ("FONTSIZE",    (0,0), (-1,-1), 8),
                        ("ALIGN",       (0,0), (-1,-1), "CENTER"),
                        ("ALIGN",       (1,1), (2,-1), "LEFT"),
                        ("GRID",        (0,0), (-1,-1), 0.4, colors.HexColor("#CCCCCC")),
                        ("FONTNAME",    (-1,1), (-1,-1), "Helvetica-Bold"),
                        ("TEXTCOLOR",   (-1,1), (-1,-1), RED),
                        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, LGREY]),
                        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
                        ("TOPPADDING",  (0,0), (-1,-1), 4),
                        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                    ]
                    tbl_obj.setStyle(TableStyle(tbl_style))

                    story = [
                        Paragraph("Federal University of Technology, Owerri", title_style),
                        Paragraph(f"{department} — {_exp_lvl} GPA Record", sub_style),
                        tbl_obj,
                    ]
                    doc.build(story)
                    return buf.getvalue()

                # ── Download buttons ──────────────────────────────────────────────
                with _ecol1:
                    try:
                        _excel_bytes = _export_excel()
                        st.download_button(
                            "📥 Download Excel",
                            data=_excel_bytes,
                            file_name=f"GPA_{department}_{_exp_lvl}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                        )
                    except Exception as _e:
                        st.error(f"Excel error: {_e}")

                with _ecol2:
                    try:
                        _word_bytes = _export_word()
                        st.download_button(
                            "📥 Download Word",
                            data=_word_bytes,
                            file_name=f"GPA_{department}_{_exp_lvl}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    except Exception as _e:
                        st.error(f"Word error: {_e}")

                with _ecol3:
                    try:
                        _pdf_bytes = _export_pdf()
                        st.download_button(
                            "📥 Download PDF",
                            data=_pdf_bytes,
                            file_name=f"GPA_{department}_{_exp_lvl}.pdf",
                            mime="application/pdf",
                            use_container_width=True,
                        )
                    except Exception as _e:
                        st.error(f"PDF error: {_e}")

        # ── TAB 6 — Department Statistics ────────────────────────────────────────
        with tab6:
            import requests as _req
            import pandas as _pd
            from io import BytesIO as _BytesIO
            import plotly.express as _px

            st.markdown(f"### 📈 Department Statistics — {department}")
            st.caption(
                "Reads attendance archives from LAVA. "
                "Select level, session and semester to load records."
            )

            # ── LAVA connection ───────────────────────────────────────────────────
            try:
                _lava_token = st.secrets["GITHUB_PAT"]
                _lava_repo  = f"{st.secrets['LAVA_OWNER']}/{st.secrets['LAVA_REPO']}"
                _lava_hdrs  = {
                    "Authorization": f"token {_lava_token}",
                    "Accept": "application/vnd.github.v3+json",
                }
            except Exception:
                st.error("LAVA secrets not configured. Add LAVA_OWNER and LAVA_REPO to secrets.")
                st.stop()

            @st.cache_data(ttl=120, show_spinner=False)
            def _gh_list_stats(path, _repo, _token):
                url  = f"https://api.github.com/repos/{_repo}/contents/{path}"
                hdrs = {"Authorization": f"token {_token}",
                        "Accept": "application/vnd.github.v3+json"}
                r = _req.get(url, headers=hdrs, timeout=10)
                if r.status_code != 200:
                    return []
                d = r.json()
                return d if isinstance(d, list) else []

            @st.cache_data(ttl=300, show_spinner=False)
            def _fetch_csv_stats(download_url, _token):
                hdrs = {"Authorization": f"token {_token}",
                        "Accept": "application/vnd.github.v3+json"}
                return _req.get(download_url, headers=hdrs, timeout=10).content

            # ── Filters ───────────────────────────────────────────────────────────
            dept_abbr       = get_dept_abbreviation(department)
            school_abbr_val = get_school_abbr(school)
            levels          = get_levels(department, school)

            # Available sessions from semester history + active
            _all_sessions = get_available_sessions()

            if not _all_sessions:
                st.info(
                    "No semesters have been started yet. "
                    "Ask ICT to start a semester before loading statistics."
                )
                st.stop()

            fc1, fc2, fc3 = st.columns([1, 1, 1])
            with fc1:
                sel_level = st.selectbox("Level", levels, key="stats_level_sel")
            with fc2:
                sel_session = st.selectbox(
                    "Academic Session",
                    _all_sessions,
                    key="stats_session_sel",
                    help="Sessions are created when ICT starts a semester.",
                )
            with fc3:
                sel_sem_choice = st.selectbox(
                    "Semester",
                    ["First Semester", "Second Semester", "Both Semesters"],
                    key="stats_sem_choice_sel",
                )

            file_prefix = f"{school_abbr_val}{dept_abbr}{sel_level}"

            if st.button("🔍 Load Statistics", type="primary", key="stats_load"):
                for k in ["stats_master", "stats_matched_files"]:
                    st.session_state.pop(k, None)
                st.session_state["stats_loaded"]      = True
                st.session_state["stats_level"]       = sel_level
                st.session_state["stats_prefix"]      = file_prefix
                st.session_state["stats_session"]     = sel_session
                st.session_state["stats_sem_choice"]  = sel_sem_choice

            if not st.session_state.get("stats_loaded"):
                st.info("Choose a level, session and semester then click **Load Statistics**.")
                st.stop()

            # ── Resolve which semester folders to scan ─────────────────────────────
            # Guard: if any required key is missing, reset and prompt reload
            _required_stats_keys = ["stats_session", "stats_sem_choice", "stats_level", "stats_prefix"]
            if any(k not in st.session_state for k in _required_stats_keys):
                for k in ["stats_loaded", "stats_master", "stats_matched_files"] + _required_stats_keys:
                    st.session_state.pop(k, None)
                st.info("Please choose your filters and click **Load Statistics**.")
                st.stop()

            _session_str = st.session_state["stats_session"]
            _sem_choice  = st.session_state["stats_sem_choice"]
            sel_level    = st.session_state["stats_level"]
            file_prefix  = st.session_state["stats_prefix"]
            _sem_records = get_semesters_for_session(_session_str)

            # Map "First Semester" / "Second Semester" / "Both Semesters" to folder names
            def _sem_folder(sem_rec: dict) -> str:
                return sem_rec.get("name", "").replace(" ", "")

            if _sem_choice == "Both Semesters":
                _target_sems = _sem_records
            else:
                _target_sems = [
                    s for s in _sem_records
                    if s.get("name", "").lower().startswith(_sem_choice.split()[0].lower())
                ]

            if not _target_sems:
                st.warning(
                    f"No semester records found for **{_session_str}** / **{_sem_choice}**. "
                    "ICT may not have started that semester yet."
                )
                st.stop()

            # ── Collect matching CSVs from LAVA ───────────────────────────────────
            if "stats_master" not in st.session_state:
                matched_files = []
                with st.spinner("Scanning LAVA archives..."):
                    for _sem_rec in _target_sems:
                        _sem_folder_name = _sem_folder(_sem_rec)
                        _base_path = f"attendances/{_session_str}/{_sem_folder_name}"
                        # List date sub-folders
                        _date_folders = [
                            f["name"] for f in _gh_list_stats(_base_path, _lava_repo, _lava_token)
                            if f.get("type") == "dir"
                        ]
                        progress = st.progress(0, text=f"Scanning {_sem_rec.get('name','')}...")
                        for _di, _date_folder in enumerate(sorted(_date_folders)):
                            progress.progress(
                                (_di + 1) / max(len(_date_folders), 1),
                                text=f"Scanning {_date_folder}..."
                            )
                            for f in _gh_list_stats(f"{_base_path}/{_date_folder}", _lava_repo, _lava_token):
                                name = f.get("name", "")
                                if not name.endswith(".csv"):
                                    continue
                                if not name.startswith(file_prefix + "_"):
                                    continue
                                # Parse: {ABBR}{LEVEL}_{SESSION}+{SEMESTER}_{COURSE}_{DATE}.csv
                                try:
                                    base   = name.replace(".csv", "")
                                    parts  = base.split("_")
                                    # parts[0] = ABBRLEVEL
                                    # parts[1] = SESSION+SEMESTER
                                    # parts[2] = COURSECODE
                                    # parts[3] = DATE
                                    course_code = parts[2]
                                    file_date   = parts[3]
                                except (IndexError, ValueError):
                                    course_code = "UNKNOWN"
                                    file_date   = _date_folder
                                matched_files.append({
                                    "name":         name,
                                    "download_url": f.get("download_url", ""),
                                    "date":         file_date,
                                    "course_code":  course_code,
                                    "semester":     _sem_rec.get("name", ""),
                                })
                        progress.empty()

                if not matched_files:
                    st.warning(
                        f"No records found for **{file_prefix}** in "
                        f"**{_session_str}** / **{_sem_choice}**."
                    )
                    st.stop()

                # Load all CSVs into master DataFrame
                all_rows = []
                prog2 = st.progress(0, text="Loading records...")
                for i, mf in enumerate(matched_files):
                    prog2.progress((i + 1) / len(matched_files), text=f"Loading {mf['name']}...")
                    try:
                        raw = _fetch_csv_stats(mf["download_url"], _lava_token)
                        df  = _pd.read_csv(_BytesIO(raw))
                        df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
                        mat_col = next((c for c in df.columns if "matric" in c), None)
                        if mat_col is None:
                            continue
                        if mat_col != "matric_number":
                            df = df.rename(columns={mat_col: "matric_number"})
                        df["course_code"] = mf["course_code"]
                        df["date"]        = mf["date"]
                        df["semester"]    = mf["semester"]
                        all_rows.append(df)
                    except Exception:
                        continue
                prog2.empty()

                if not all_rows:
                    st.error("Could not parse any attendance records.")
                    st.stop()

                master = _pd.concat(all_rows, ignore_index=True)
                master["matric_number"] = master["matric_number"].astype(str).str.strip()

                if "surname" in master.columns and "other_names" in master.columns:
                    master["full_name"] = (
                        master["surname"].fillna("").str.strip().str.upper()
                        + " "
                        + master["other_names"].fillna("").str.strip().str.title()
                    ).str.strip()
                else:
                    master["full_name"] = master["matric_number"]

                st.session_state["stats_master"]        = master
                st.session_state["stats_matched_files"] = matched_files

            master        = st.session_state["stats_master"]
            matched_files = st.session_state["stats_matched_files"]
            course_list   = sorted(master["course_code"].unique())

            # ── Info banner ────────────────────────────────────────────────────────
            st.markdown(f"""<div class="info-card">
                <b>Level:</b> {sel_level} &nbsp;|&nbsp;
                <b>Session:</b> {_session_str} &nbsp;|&nbsp;
                <b>Semester:</b> {_sem_choice} &nbsp;|&nbsp;
                <b>Records:</b> {len(matched_files)} files &nbsp;|&nbsp;
                <b>Entries:</b> {len(master)}
            </div>""", unsafe_allow_html=True)

            # ── Overview stats ─────────────────────────────────────────────────────
            st.markdown("### Overview")
            ov1, ov2, ov3 = st.columns(3)
            for col, num, lbl in [
                (ov1, master["matric_number"].nunique(), "Unique Students"),
                (ov2, master["course_code"].nunique(),   "Courses"),
                (ov3, len(master),                       "Total Entries"),
            ]:
                col.markdown(
                    f'<div class="stat-box"><div class="num">{num}</div>'
                    f'<div class="lbl">{lbl}</div></div>',
                    unsafe_allow_html=True,
                )

            st.markdown("<br>", unsafe_allow_html=True)

            # ── Course summary table ───────────────────────────────────────────────
            st.markdown("### Attendance per Course")
            course_summary = (
                master.groupby("course_code")
                .agg(sessions=("date", "nunique"), total_entries=("matric_number", "count"))
                .reset_index()
                .rename(columns={
                    "course_code":   "Course Code",
                    "sessions":      "Sessions Held",
                    "total_entries": "Total Entries",
                })
                .sort_values("Course Code")
            )
            st.dataframe(course_summary, use_container_width=True, hide_index=True)

            st.divider()

            # ── Student drill-down ─────────────────────────────────────────────────
            st.markdown("### 👤 Student Drill-Down")

            name_map = (
                master.dropna(subset=["matric_number"])
                .drop_duplicates(subset=["matric_number"])
                .set_index("matric_number")["full_name"]
                .to_dict()
            )

            student_opts = {
                f"{mat}  —  {name_map.get(mat, '')}": mat
                for mat in sorted(name_map.keys())
            }

            sel_label  = st.selectbox("Select student", list(student_opts.keys()), key="stats_stu")
            sel_matric = student_opts[sel_label]
            sel_name   = name_map.get(sel_matric, sel_matric)

            student_rows     = master[master["matric_number"] == sel_matric]
            total_attended   = len(student_rows)
            courses_attended = student_rows["course_code"].nunique()

            st.markdown(f"""<div class="info-card">
                <b>{sel_name}</b> &nbsp;|&nbsp;
                Matric: <b>{sel_matric}</b> &nbsp;|&nbsp;
                Total entries: <b>{total_attended}</b> &nbsp;|&nbsp;
                Courses attended: <b>{courses_attended}</b>
            </div>""", unsafe_allow_html=True)

            # Per-student attendance summary (for export)
            stu_export = (
                student_rows.groupby("course_code")
                .size()
                .reset_index(name="Times Attended")
                .rename(columns={"course_code": "Course Code"})
            )
            all_c_df   = _pd.DataFrame({"Course Code": course_list})
            stu_export = all_c_df.merge(stu_export, on="Course Code", how="left").fillna(0)
            stu_export["Times Attended"] = stu_export["Times Attended"].astype(int)
            stu_export["Sessions Held"]  = stu_export["Course Code"].map(
                master.groupby("course_code")["date"].nunique()
            ).fillna(0).astype(int)
            stu_export = stu_export[["Course Code", "Sessions Held", "Times Attended"]]

            # Time / lateness columns
            _has_time    = "time" in master.columns
            _has_started = "session_started" in master.columns

            def _parse_hms(s):
                try:
                    parts = str(s).strip().split(":")
                    return int(parts[0]) * 60 + int(parts[1])
                except Exception:
                    return None

            log_cols_src  = ["course_code", "date", "semester"]
            log_cols_disp = {"course_code": "Course Code", "date": "Date", "semester": "Semester"}

            if _has_time:
                log_cols_src.append("time")
                log_cols_disp["time"] = "Sign-in Time"
            if _has_started:
                log_cols_src.append("session_started")
                log_cols_disp["session_started"] = "Class Started"

            # Only keep columns that actually exist
            log_cols_src = [c for c in log_cols_src if c in student_rows.columns]

            stu_log = (
                student_rows[log_cols_src]
                .sort_values(["date", "course_code"])
                .rename(columns=log_cols_disp)
                .reset_index(drop=True)
            )

            if _has_time and _has_started:
                def _late_status(row):
                    signin  = _parse_hms(row.get("Sign-in Time", ""))
                    started = _parse_hms(row.get("Class Started", ""))
                    if signin is None or started is None:
                        return "—"
                    diff = signin - started
                    if diff <= 0:
                        return "✅ On Time"
                    elif diff <= 15:
                        return f"⚠️ Late ({diff}m)"
                    else:
                        return f"🔴 Very Late ({diff}m)"
                stu_log["Status"] = stu_log.apply(_late_status, axis=1)

            # ── Load GPA data for selected student ────────────────────────────────
            _gpa_records = load_student_gpa(sel_matric)
            _cgpa        = compute_cgpa(_gpa_records) if _gpa_records else None

            # Filter GPA records to selected session / semester choice
            def _gpa_matches(rec: dict) -> bool:
                rec_session = rec.get("session", "").replace("/", "-")
                rec_name    = rec.get("semester", "")
                if rec_session != _session_str:
                    return False
                if _sem_choice == "Both Semesters":
                    return True
                return rec_name.lower().startswith(_sem_choice.split()[0].lower())

            _period_gpa = [r for r in _gpa_records if _gpa_matches(r)]

            # ── Charts ────────────────────────────────────────────────────────────
            chart_view = st.radio(
                "Chart view",
                ["By Course Code", "By Date"],
                horizontal=True,
                key="stats_chart_view",
            )

            import hashlib as _hashlib

            def _course_colour(code: str) -> str:
                h   = int(_hashlib.md5(code.encode()).hexdigest()[:6], 16)
                hue = h % 360
                return f"hsl({hue}, 72%, 55%)"

            def _chart_height(n_bars: int) -> int:
                return max(260, min(420, 180 + n_bars * 36))

            CHART_LAYOUT_BASE = dict(
                plot_bgcolor="rgba(0,0,0,0)",
                paper_bgcolor="rgba(0,0,0,0)",
                font_color="#e8e4dc",
                title_font_size=13,
                showlegend=False,
                margin=dict(t=44, b=36, l=8, r=8),
                yaxis=dict(tickformat="d", gridcolor="rgba(255,255,255,0.07)", dtick=1),
                xaxis=dict(gridcolor="rgba(0,0,0,0)", tickangle=-30, tickfont=dict(size=11)),
            )

            if chart_view == "By Course Code":
                sc = (
                    student_rows.groupby("course_code")
                    .size()
                    .reset_index(name="times_attended")
                )
                all_courses_df = _pd.DataFrame({"course_code": course_list})
                sc = all_courses_df.merge(sc, on="course_code", how="left").fillna(0)
                sc["times_attended"] = sc["times_attended"].astype(int)
                sc["colour"] = sc["course_code"].apply(_course_colour)
                fig = _px.bar(
                    sc, x="course_code", y="times_attended", text="times_attended",
                    labels={"course_code": "Course Code", "times_attended": "Times Attended"},
                    title=f"{sel_name} — Attendance by Course",
                    color="course_code",
                    color_discrete_map={row["course_code"]: row["colour"] for _, row in sc.iterrows()},
                )
                fig.update_traces(textposition="outside", marker_line_width=0, textfont_size=11)
                fig.update_layout(**CHART_LAYOUT_BASE, height=_chart_height(len(sc)))
                st.plotly_chart(fig, use_container_width=True)
            else:
                sd = (
                    student_rows.groupby("date")
                    .size()
                    .reset_index(name="classes_attended")
                    .sort_values("date")
                )
                fig = _px.bar(
                    sd, x="date", y="classes_attended", text="classes_attended",
                    labels={"date": "Date", "classes_attended": "Classes Attended"},
                    title=f"{sel_name} — Attendance by Date",
                    color_discrete_sequence=["rgba(41,128,185,0.85)"],
                )
                fig.update_traces(textposition="outside", marker_line_width=0, textfont_size=11)
                fig.update_layout(**CHART_LAYOUT_BASE, height=_chart_height(len(sd)))
                st.plotly_chart(fig, use_container_width=True)

            st.markdown("#### Attendance Log")
            breakdown = stu_log.copy()
            breakdown.index = range(1, len(breakdown) + 1)
            st.dataframe(breakdown, use_container_width=True)

            # ── Full matrix ────────────────────────────────────────────────────────
            st.divider()
            st.markdown("### 📋 Full Attendance Matrix")
            st.caption("Rows = students · Columns = course codes · Values = times attended")
            pivot = (
                master.groupby(["matric_number", "course_code"])
                .size()
                .unstack(fill_value=0)
                .reset_index()
            )
            pivot.insert(1, "Name", pivot["matric_number"].map(name_map))
            pivot = pivot.rename(columns={"matric_number": "Matric No."})
            course_cols_p = [c for c in pivot.columns if c not in ("Matric No.", "Name")]
            pivot["Total"] = pivot[course_cols_p].sum(axis=1)
            pivot = pivot.sort_values("Total", ascending=False).reset_index(drop=True)
            pivot.index += 1
            st.dataframe(pivot, use_container_width=True)

            # ── Student report export ──────────────────────────────────────────────
            st.divider()
            st.markdown("### 📤 Export Student Report")
            st.caption(
                f"Downloads a report for **{sel_name}** ({sel_matric}) — "
                "suitable for sending to parents/guardians."
            )

            for k in ["export_remark", "export_fmt", "remark_saved"]:
                if k not in st.session_state:
                    st.session_state[k] = "" if k == "export_remark" else (None if k == "export_fmt" else False)

            ex1, ex2, ex3 = st.columns(3)
            for col, label, fmt in [
                (ex1, "📊 Excel", "xlsx"),
                (ex2, "📝 Word",  "docx"),
                (ex3, "📑 PDF",   "pdf"),
            ]:
                with col:
                    if st.button(label, key=f"exp_btn_{fmt}", use_container_width=True):
                        st.session_state["export_fmt"]    = fmt
                        st.session_state["export_remark"] = ""
                        st.session_state["remark_saved"]  = False

            if st.session_state["export_fmt"]:
                fmt = st.session_state["export_fmt"]
                st.markdown(f"#### ✍️ Remark for {sel_name}'s report")
                with st.form("remark_form"):
                    remark_input = st.text_area(
                        "Advisor's remark to parents (optional)",
                        value=st.session_state["export_remark"],
                        placeholder="e.g. This student has shown consistent attendance.",
                        height=120,
                    )
                    save_remark = st.form_submit_button("💾 Save Remark & Generate Download")
                if save_remark:
                    st.session_state["export_remark"] = remark_input
                    st.session_state["remark_saved"]  = True
                remark = st.session_state["export_remark"]

                _safe_name  = sel_name.replace(' ', '-').replace('/', '-')
                base_name   = f"{dept_abbr}{sel_level}_{sel_matric}_{_safe_name}_{_session_str}_{_sem_choice.replace(' ','')}_attendance"
                advisor_sig = adv.get("username", "Course Advisor")

                if not st.session_state.get("remark_saved"):
                    st.info("Type your remark above (or leave blank) and tap **Save Remark & Generate Download**.")

                # ── Shared GPA block builder ───────────────────────────────────────
                # Used by all three export formats
                _period_label = f"{_session_str} — {_sem_choice}"

                # ── Excel ──────────────────────────────────────────────────────────
                def _to_excel():
                    import io as _io
                    from openpyxl import Workbook
                    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

                    wb  = Workbook()
                    ws  = wb.active
                    ws.title = "Attendance Report"
                    thin   = Side(style="thin", color="CCCCCC")
                    border = Border(left=thin, right=thin, top=thin, bottom=thin)
                    RED    = "7B0000"
                    hdr_fill = PatternFill("solid", fgColor=RED)
                    alt_fill = PatternFill("solid", fgColor="FDF2F2")

                    # Letterhead
                    ws.merge_cells("A1:E1")
                    ws["A1"] = "FEDERAL UNIVERSITY OF TECHNOLOGY, OWERRI"
                    ws["A1"].font      = Font(name="Arial", bold=True, size=13, color=RED)
                    ws["A1"].alignment = Alignment(horizontal="center")

                    ws.merge_cells("A2:E2")
                    ws["A2"] = f"Student Attendance Report  |  {department}  |  Level {sel_level}"
                    ws["A2"].font      = Font(name="Arial", size=10, italic=True)
                    ws["A2"].alignment = Alignment(horizontal="center")

                    ws.merge_cells("A3:E3")
                    ws["A3"] = f"Session: {_session_str}   Semester: {_sem_choice}"
                    ws["A3"].font      = Font(name="Arial", size=10)
                    ws["A3"].alignment = Alignment(horizontal="center")

                    # Student info
                    ws["A5"] = "Student Name"; ws["B5"] = sel_name
                    ws["A6"] = "Matric No.";   ws["B6"] = sel_matric
                    ws["A7"] = "Total Entries"; ws["B7"] = total_attended
                    for row in [5, 6, 7]:
                        ws.cell(row, 1).font = Font(name="Arial", bold=True, size=10)
                        ws.cell(row, 2).font = Font(name="Arial", size=10)

                    # Attendance summary table — row 9
                    for ci, h in enumerate(stu_export.columns, 1):
                        c = ws.cell(9, ci, value=h)
                        c.font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                        c.fill = hdr_fill; c.alignment = Alignment(horizontal="center"); c.border = border

                    for ri, row in enumerate(stu_export.itertuples(index=False), start=10):
                        for ci, val in enumerate(row, 1):
                            c = ws.cell(ri, ci, value=val)
                            c.font = Font(name="Arial", size=10)
                            c.alignment = Alignment(horizontal="center" if ci > 1 else "left")
                            c.border = border
                            if ri % 2 == 0: c.fill = alt_fill

                    next_row = 10 + len(stu_export) + 2

                    # GPA section
                    if _gpa_records:
                        ws.cell(next_row, 1).value = "GPA / CGPA Summary"
                        ws.cell(next_row, 1).font  = Font(name="Arial", bold=True, size=11, color=RED)
                        next_row += 1
                        # Overall CGPA
                        ws.cell(next_row, 1).value = "Overall CGPA (all semesters)"
                        ws.cell(next_row, 1).font  = Font(name="Arial", bold=True, size=10)
                        ws.cell(next_row, 2).value = f"{_cgpa:.2f}" if _cgpa is not None else "—"
                        ws.cell(next_row, 2).font  = Font(name="Arial", size=10)
                        next_row += 1
                        # Period GPA rows
                        if _period_gpa:
                            for _gr in _period_gpa:
                                ws.cell(next_row, 1).value = f"GPA — {_gr.get('semester','')}"
                                ws.cell(next_row, 1).font  = Font(name="Arial", bold=True, size=10)
                                ws.cell(next_row, 2).value = f"{_gr['gpa']:.2f}"
                                ws.cell(next_row, 2).font  = Font(name="Arial", size=10)
                                next_row += 1
                    next_row += 1

                    # Remark
                    ws.cell(next_row, 1).value = "Advisor's Remark:"
                    ws.cell(next_row, 1).font  = Font(name="Arial", bold=True, size=10)
                    next_row += 1
                    ws.merge_cells(f"A{next_row}:E{next_row+2}")
                    ws.cell(next_row, 1).value     = remark if remark.strip() else "—"
                    ws.cell(next_row, 1).font      = Font(name="Arial", size=10, italic=True)
                    ws.cell(next_row, 1).alignment = Alignment(wrap_text=True, vertical="top")
                    next_row += 3
                    ws.cell(next_row, 1).value = f"Signed: {advisor_sig}   Date: {_period_label}"
                    ws.cell(next_row, 1).font  = Font(name="Arial", size=9, italic=True)

                    ws.column_dimensions["A"].width = 28
                    ws.column_dimensions["B"].width = 20
                    ws.column_dimensions["C"].width = 16

                    buf = _io.BytesIO(); wb.save(buf); return buf.getvalue()

                # ── Word ───────────────────────────────────────────────────────────
                def _to_word():
                    import io as _io
                    from docx import Document
                    from docx.shared import Pt, RGBColor, Cm
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    def _shd(cell, hx):
                        tc=cell._tc; p=tc.get_or_add_tcPr()
                        s=OxmlElement("w:shd")
                        s.set(qn("w:fill"),hx); s.set(qn("w:color"),"auto"); s.set(qn("w:val"),"clear")
                        p.append(s)

                    doc = Document()
                    for sec in doc.sections:
                        sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
                        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

                    # Letterhead
                    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=p.add_run("FEDERAL UNIVERSITY OF TECHNOLOGY, OWERRI")
                    r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0x7B,0,0)
                    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r2=p2.add_run(f"Student Attendance Report  ·  {department}  ·  Level {sel_level}")
                    r2.font.size=Pt(10); r2.italic=True
                    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    p3.add_run(f"Session: {_session_str}   ·   Semester: {_sem_choice}").font.size=Pt(10)
                    doc.add_paragraph()

                    # Student info box
                    info_tbl = doc.add_table(rows=2, cols=4); info_tbl.style = "Table Grid"
                    labels = ["Student Name", sel_name, "Matric No.", sel_matric,
                              "Total Entries", str(total_attended), "Courses", str(courses_attended)]
                    for i in range(2):
                        for j in range(4):
                            cell = info_tbl.rows[i].cells[j]
                            cell.text = labels[i*4+j]
                            run = cell.paragraphs[0].runs[0]
                            run.font.size = Pt(10)
                            run.bold = (j % 2 == 0)
                    doc.add_paragraph()

                    # Attendance summary table
                    doc.add_paragraph().add_run("Attendance Summary").bold = True
                    tbl = doc.add_table(rows=1, cols=len(stu_export.columns))
                    tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
                    for i, h in enumerate(stu_export.columns):
                        cell=tbl.rows[0].cells[i]; cell.text=str(h)
                        run=cell.paragraphs[0].runs[0]
                        run.bold=True; run.font.size=Pt(9); run.font.color.rgb=RGBColor(255,255,255)
                        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; _shd(cell,"7B0000")
                    for ri, row in enumerate(stu_export.itertuples(index=False)):
                        cells=tbl.add_row().cells
                        for ci, val in enumerate(row):
                            cells[ci].text=str(val)
                            cells[ci].paragraphs[0].runs[0].font.size=Pt(9)
                            cells[ci].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
                            if ri%2==0: _shd(cells[ci],"FDF2F2")
                    doc.add_paragraph()

                    # GPA section
                    if _gpa_records:
                        gpa_head = doc.add_paragraph()
                        gpa_head.add_run("GPA / CGPA Summary").bold = True
                        gpa_head.runs[0].font.size = Pt(11)
                        gpa_head.runs[0].font.color.rgb = RGBColor(0x7B,0,0)
                        # CGPA row
                        gpa_tbl = doc.add_table(rows=1, cols=2); gpa_tbl.style="Table Grid"
                        _gh = gpa_tbl.rows[0].cells
                        _gh[0].text = "Overall CGPA (all semesters)"
                        _gh[0].paragraphs[0].runs[0].bold = True
                        _gh[0].paragraphs[0].runs[0].font.size = Pt(10)
                        _gh[1].text = f"{_cgpa:.2f}" if _cgpa is not None else "—"
                        _gh[1].paragraphs[0].runs[0].font.size = Pt(10)
                        _shd(_gh[0], "F5F5F5")
                        for _gr in _period_gpa:
                            cells2 = gpa_tbl.add_row().cells
                            cells2[0].text = f"GPA — {_gr.get('semester','')}"
                            cells2[0].paragraphs[0].runs[0].bold = True
                            cells2[0].paragraphs[0].runs[0].font.size = Pt(10)
                            cells2[1].text = f"{_gr['gpa']:.2f}"
                            cells2[1].paragraphs[0].runs[0].font.size = Pt(10)
                        doc.add_paragraph()

                    # Remark
                    doc.add_paragraph()
                    rp=doc.add_paragraph(); rp.add_run("Advisor's Remark:").bold=True
                    rp.runs[0].font.size=Pt(10)
                    rb=doc.add_paragraph(remark if remark.strip() else "—")
                    rb.runs[0].font.size=Pt(10); rb.runs[0].italic=True
                    doc.add_paragraph()
                    sig=doc.add_paragraph(f"Signed: {advisor_sig}          Period: {_period_label}")
                    sig.runs[0].font.size=Pt(9); sig.runs[0].italic=True

                    buf=_io.BytesIO(); doc.save(buf); return buf.getvalue()

                # ── PDF ────────────────────────────────────────────────────────────
                def _to_pdf():
                    import io as _io
                    from reportlab.lib.pagesizes import A4
                    from reportlab.lib import colors
                    from reportlab.lib.units import cm
                    from reportlab.platypus import (
                        SimpleDocTemplate, Table, TableStyle,
                        Paragraph, Spacer,
                    )
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.enums import TA_CENTER, TA_LEFT

                    buf = _io.BytesIO()
                    doc = SimpleDocTemplate(buf, pagesize=A4,
                                            topMargin=1.5*cm, bottomMargin=1.5*cm,
                                            leftMargin=1.8*cm, rightMargin=1.8*cm)
                    styles  = getSampleStyleSheet()
                    RED     = colors.HexColor("#7B0000")
                    LTRED   = colors.HexColor("#FDF2F2")
                    GREY    = colors.HexColor("#555555")

                    h1 = ParagraphStyle("h1", parent=styles["Normal"],
                        fontSize=13, fontName="Helvetica-Bold", textColor=RED,
                        alignment=TA_CENTER, spaceAfter=3)
                    sub = ParagraphStyle("sub", parent=styles["Normal"],
                        fontSize=9, alignment=TA_CENTER, textColor=GREY, spaceAfter=2)
                    label_s = ParagraphStyle("lbl", parent=styles["Normal"],
                        fontSize=10, fontName="Helvetica-Bold", spaceAfter=2)
                    remark_s = ParagraphStyle("rmk", parent=styles["Normal"],
                        fontSize=10, leftIndent=10, textColor=GREY,
                        borderPad=6, spaceAfter=6)

                    story = [
                        Paragraph("FEDERAL UNIVERSITY OF TECHNOLOGY, OWERRI", h1),
                        Paragraph(f"Student Attendance Report  ·  {department}  ·  Level {sel_level}", sub),
                        Paragraph(f"Session: {_session_str}   ·   Semester: {_sem_choice}", sub),
                        Spacer(1, 0.3*cm),
                    ]

                    # Student info table
                    info_data = [
                        ["Student Name", sel_name, "Matric No.", sel_matric],
                        ["Total Entries", str(total_attended), "Courses", str(courses_attended)],
                    ]
                    info_tbl = Table(info_data, colWidths=[3.5*cm, 5*cm, 3*cm, 4*cm])
                    info_tbl.setStyle(TableStyle([
                        ("FONTNAME",   (0,0),(-1,-1), "Helvetica"),
                        ("FONTNAME",   (0,0),(0,-1),  "Helvetica-Bold"),
                        ("FONTNAME",   (2,0),(2,-1),  "Helvetica-Bold"),
                        ("FONTSIZE",   (0,0),(-1,-1), 9),
                        ("GRID",       (0,0),(-1,-1), 0.4, colors.HexColor("#DDDDDD")),
                        ("BACKGROUND", (0,0),(0,-1),  colors.HexColor("#F5F5F5")),
                        ("BACKGROUND", (2,0),(2,-1),  colors.HexColor("#F5F5F5")),
                    ]))
                    story += [info_tbl, Spacer(1, 0.4*cm)]

                    # Attendance summary table
                    story.append(Paragraph("Attendance Summary", label_s))
                    s_headers = list(stu_export.columns)
                    s_data    = [s_headers] + [list(r) for r in stu_export.itertuples(index=False)]
                    s_tbl = Table(s_data, colWidths=[5*cm, 4*cm, 4*cm], repeatRows=1)
                    s_tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0,0),(-1,0), RED),
                        ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
                        ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
                        ("FONTSIZE",   (0,0),(-1,-1), 9),
                        ("ALIGN",      (0,0),(-1,-1), "CENTER"),
                        ("ALIGN",      (0,1),(0,-1),  "LEFT"),
                        ("GRID",       (0,0),(-1,-1), 0.4, colors.HexColor("#DDDDDD")),
                        *[("BACKGROUND",(0,i),(-1,i), LTRED) for i in range(2, len(s_data), 2)],
                    ]))
                    story += [s_tbl, Spacer(1, 0.4*cm)]

                    # GPA section
                    if _gpa_records:
                        story.append(Paragraph("GPA / CGPA Summary", label_s))
                        gpa_data = [["Semester", "GPA"]]
                        if _cgpa is not None:
                            gpa_data.append(["Overall CGPA (all semesters)", f"{_cgpa:.2f}"])
                        for _gr in _period_gpa:
                            gpa_data.append([f"GPA — {_gr.get('semester','')}", f"{_gr['gpa']:.2f}"])
                        g_tbl = Table(gpa_data, colWidths=[10*cm, 4*cm], repeatRows=1)
                        g_tbl.setStyle(TableStyle([
                            ("BACKGROUND", (0,0),(-1,0), RED),
                            ("TEXTCOLOR",  (0,0),(-1,0), colors.white),
                            ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
                            ("FONTNAME",   (0,1),(0,-1), "Helvetica-Bold"),
                            ("FONTSIZE",   (0,0),(-1,-1), 9),
                            ("ALIGN",      (0,0),(-1,-1), "LEFT"),
                            ("ALIGN",      (1,0),(-1,-1), "CENTER"),
                            ("GRID",       (0,0),(-1,-1), 0.4, colors.HexColor("#DDDDDD")),
                            *[("BACKGROUND",(0,i),(-1,i), LTRED) for i in range(2, len(gpa_data), 2)],
                        ]))
                        story += [g_tbl, Spacer(1, 0.4*cm)]

                    # Remark
                    story.append(Paragraph("Advisor's Remark:", label_s))
                    story.append(Paragraph(remark if remark.strip() else "—", remark_s))
                    story.append(Spacer(1, 0.3*cm))
                    story.append(Paragraph(
                        f"Signed: <b>{advisor_sig}</b>          Period: {_period_label}",
                        ParagraphStyle("sig", parent=styles["Normal"], fontSize=9,
                                       textColor=GREY, alignment=TA_LEFT)
                    ))

                    doc.build(story)
                    return buf.getvalue()

                # ── Render download button ─────────────────────────────────────────
                try:
                    if fmt == "xlsx":
                        data_bytes = _to_excel()
                        mime       = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    elif fmt == "docx":
                        data_bytes = _to_word()
                        mime       = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    else:
                        data_bytes = _to_pdf()
                        mime       = "application/pdf"

                    st.download_button(
                        f"⬇️ Download {fmt.upper()} for {sel_name}",
                        data_bytes,
                        file_name=f"{base_name}.{fmt}",
                        mime=mime,
                        type="primary",
                        use_container_width=True,
                        key=f"dl_{fmt}_{sel_matric}",
                    )
                except Exception as e:
                    st.error(f"Export failed: {e}")



except Exception as _err:
    if type(_err).__name__ in ("StopException", "RerunException"):
        raise
    _show_error(_err)

